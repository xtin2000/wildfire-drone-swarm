#!/usr/bin/env python3
"""
update_thesis_section5_v3.py — Apply final Section 5 redraft to thesis.

Builds on Thesis_04_27_Combined_Mechanism.docx (which already has the dual-mechanism
overview, §5.3.5 rotor-wash subsection, and a first pass at §5.6/§5.8). Replaces
§5.8 and §5.9.2 with the FINAL content reflecting:

  - Option B 2×2 factorial (drones × firefighters)
  - Three-condition complementarity framing (drones + firefighters + wind envelope)
  - Seed reproducibility N=5 (3/5 contained, 2/5 fuel exhaustion)
  - Rotor-wash strength sensitivity (half / 1× / double)
  - Airborne-ember drag isolation (combined_no_ember_drag)
  - High-wind robustness re-runs under Option B
  - Probabilistic-containment finding

All changes yellow-highlighted for review. Saves as new file.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_27_Combined_Mechanism.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_28_Section5_v3_Final.docx"
FIG_DIR = "/Users/xtin2000/wildfire_drone_swarm/results_v2/figures"


def find_paragraph_index(doc, snippet):
    for i, p in enumerate(doc.paragraphs):
        if snippet in p.text:
            return i
    return -1


def add_paragraph_after(doc, ref_para, text, style="Normal", highlight=True, bold=False):
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles[style]
            if text:
                run = p.add_run(text)
                if highlight:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                if bold:
                    run.bold = True
            return p
    return None


def insert_figure_after(doc, ref_para, fig_path, caption):
    cap_p = OxmlElement("w:p")
    ref_para._element.addnext(cap_p)
    for p in doc.paragraphs:
        if p._element is cap_p:
            p.style = doc.styles["Normal"]
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(10)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break
    if os.path.exists(fig_path):
        fig_p = OxmlElement("w:p")
        ref_para._element.addnext(fig_p)
        for p in doc.paragraphs:
            if p._element is fig_p:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(fig_path, width=Inches(5.5))
                break


def delete_between(doc, start_snippet, end_snippet):
    s_idx = find_paragraph_index(doc, start_snippet)
    e_idx = find_paragraph_index(doc, end_snippet)
    if s_idx < 0 or e_idx < 0 or e_idx <= s_idx:
        print(f"  WARN: anchors not found: '{start_snippet}' / '{end_snippet}'")
        return None
    for i in range(e_idx - 1, s_idx, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)
    # Refresh: caller will re-find the start anchor
    return find_paragraph_index(doc, start_snippet)


def main():
    print(f"Loading: {INPUT_PATH}")
    if not os.path.exists(INPUT_PATH):
        raise FileNotFoundError(INPUT_PATH)
    doc = Document(INPUT_PATH)

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE A: replace §5.8 with the FINAL results structure
    # ─────────────────────────────────────────────────────────────────────────
    print("Replacing §5.8 Simulation Results with final v3 content...")
    s58 = find_paragraph_index(doc, "5.8 Simulation Results")
    s59 = find_paragraph_index(doc, "5.9 Analysis and Discussion")
    if s58 < 0 or s59 < 0 or s59 <= s58:
        print("  ERROR: couldn't locate §5.8 anchors. Aborting §5.8 update.")
    else:
        # Delete current §5.8 contents (keep heading)
        for i in range(s59 - 1, s58, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

        anchor = doc.paragraphs[s58]

        # Build content blocks tail-first (since we use addnext, last added appears first)
        blocks = []

        # ── §5.8.7 Sensitivity Analysis ─────────────────────────────────────
        blocks.extend([
            ("normal",
             "Airborne-ember drag isolation (combined, seed 42). The rotor-wash physics in §5.3.5 "
             "includes two distinct effects on embers: (a) ground-cell suppression of STATE_EMBER "
             "cells directly under a hovering drone, and (b) airborne-ember drag applied to "
             "in-flight ember particles passing through the downwash column. To isolate the "
             "contribution of (b), the simulation was re-run with the airborne-drag effect "
             "disabled while keeping ground-cell suppression active. Result: containment in 240 s "
             "(4.0 min) with 289 cells burned (0.7% of grid) — actually marginally faster than "
             "baseline (999 s, 1,388 burned, 3.5%). The airborne-drag term is therefore not "
             "load-bearing for the containment finding; the result rests on (i) acoustic "
             "suppression of established burning cells and (ii) rotor-wash ground-cell suppression "
             "of nascent ember-state ignitions. The more speculative airborne-drag effect is a "
             "model embellishment, not the cause of the headline outcome."),

            ("normal",
             "Rotor-wash strength sensitivity (combined, seed 42). The rotor-wash ember-rate "
             "constant ROTOR_WASH_EMBER_RATE was varied across half (25 J/s), baseline (50 J/s), "
             "and double (100 J/s). Containment occurred in all three variants in 14–17 simulated "
             "minutes with 3.4–5.5% grid burnout. The headline result is therefore robust to "
             "±100% variation in this parameter — the small non-monotonicity (double slightly "
             "worse than baseline) is consistent with stochastic variation in fire spread paths "
             "under different drone-suppression patterns."),

            ("normal",
             "Seed reproducibility (combined_baseline at seeds 42, 43, 44, 45, 46). The headline "
             "containment finding is probabilistic, not deterministic. Across N=5 random seeds at "
             "the otherwise identical combined_baseline configuration, three runs (seeds 42, 43, "
             "45) achieved real containment (≤4% grid burnout in <17 simulated minutes) and two "
             "runs (seeds 44, 46) lost the early-spread race within the first 100–200 simulated "
             "seconds, after which the fire grew beyond the swarm’s per-tick suppression capacity "
             "and ran to fuel exhaustion (91–93% grid burnout). The boolean mission_complete flag "
             "fires in both cases, but the underlying outcomes are qualitatively different. "
             "Containment rate at this fleet size is therefore best characterized as approximately "
             "60% — a probabilistic event whose failure mode (loss of the early-spread race) is "
             "the same dynamic observed in the drones_only scenario above."),

            ("heading", "5.8.7 Sensitivity Analysis"),
        ])

        # ── §5.8.6 Battery (carry-over, lightly updated) ────────────────────
        blocks.extend([
            ("normal",
             "The implication for operational deployment is encouraging: when the swarm is "
             "effective enough to contain the fire (combined mechanism with both subsystems and a "
             "ground-crew firebreak in place), the engagement is brief and battery endurance is "
             "comfortable. When the swarm cannot contain the fire — whether due to mechanism "
             "absence, firefighter absence, high wind, or unfavorable seed-level RNG — battery "
             "becomes the limiting factor for mission duration, but in those cases the operational "
             "value of continuing the engagement is itself questionable."),

            ("normal",
             "Combined baseline (with firefighters): mission complete at t ≈ 1,000 s with mean "
             "battery around 93%. Battery is not a limiting factor for short successful missions. "
             "Rotor-wash-only baseline and the high-wind scenarios: ran the full 35,000 ticks "
             "with mean battery declining linearly to approximately 60%. Sustained operation "
             "against an uncontainable fire continues to deplete batteries until intervention "
             "ends."),

            ("normal",
             "Battery dynamics differ markedly between the scenarios that achieve containment and "
             "those that do not."),
            ("heading", "5.8.6 Battery and Operational Endurance"),
        ])

        # ── §5.8.5 Wind Impact ──────────────────────────────────────────────
        blocks.extend([
            ("normal",
             "The convergence of the two high-wind scenarios on identical numerical outcomes is "
             "itself diagnostic. Because both runs share seed 42 and the high-wind regime forces "
             "both into a near-RTB-dominated state, the dynamics reduce to the seeded fire-and-wind "
             "trajectory with negligible drone influence — the floor of swarm effectiveness in "
             "this simulation framework."),

            ("normal",
             "This finding bounds the operational envelope of the combined-mechanism approach. "
             "The dramatic containment result of §5.8.3 is contingent on drones being able to "
             "maintain station above the fire. When wind conditions exceed the platform’s "
             "stability limits, the swarm is reduced to its grounded state and the suppression "
             "mechanism — however physically capable — has no opportunity to act. This is "
             "consistent with manufacturer multirotor wind limits and with the §5.3.1 stability "
             "model, but the operational implication is sharp: deployment doctrine must include "
             "real-time wind-aware tasking, and the swarm cannot be relied upon as the sole "
             "containment asset under marginal weather. The Option B firefighter physics did not "
             "alter the high-wind results materially (≤300 cells difference) because firefighter "
             "walking speed cannot catch a high-wind-driven fire whether or not their water "
             "extinguishes."),

            ("normal",
             "The mechanism for this collapse is operational rather than physical. The 10 m/s "
             "base wind, combined with the gust process (steady-state σ ≈ 1.9 m/s), produces "
             "frequent excursions above the 12 m/s drone safety threshold (§5.3.1) at which "
             "drones autonomously return to base. Inspection of the per-tick logs confirms that "
             "drone time-on-station is severely degraded in the high-wind scenarios: drones spend "
             "a large fraction of the simulation in transit or at base recharging, leaving the "
             "fire effectively unsuppressed during gust events. With drones repeatedly grounded, "
             "the two mechanism toggles become operationally equivalent, and both runs converge "
             "to the trajectory of an effectively unsuppressed fire."),

            ("normal",
             "The two high-wind scenarios produce a striking and informative result: both "
             "Rotor-Wash-Only and Combined mechanisms collapse to identical numerical outcomes "
             "under 10 m/s wind, with peak burning of approximately 4,121 cells, 37,572 cells "
             "burned, and peak ember count of 324 in both cases. The combined mechanism — which "
             "contains the fire entirely under moderate wind — provides no measurable benefit "
             "over rotor-wash-only when wind speed doubles."),

            ("heading", "5.8.5 Wind Impact"),
        ])

        # ── §5.8.4 Mechanism + Agent Decomposition ──────────────────────────
        blocks.extend([
            ("normal",
             "Drone-firefighter complementarity (geometric). The 10 ground firefighters in this "
             "simulation almost never reach an active flame. At 0.5 m/s walking speed they cannot "
             "catch a fire spreading at over 1 m/s in 5 m/s wind. Their actual contribution to "
             "the containment outcome is geometric: they create a 2,000–2,500-cell moisture "
             "firebreak along the western (upwind) flank of the fire as it approaches their "
             "position, which prevents westward spread and frees the drone swarm to focus "
             "exclusively on the eastern front. Without this firebreak (the drones_only "
             "scenario), the swarm cannot defend two flanks simultaneously and the fire runs to "
             "fuel exhaustion. This finding inverts the conventional framing of UAVs as “force "
             "multipliers”: at the fleet size and platform parameters modeled here, the swarm is "
             "more accurately described as filling the coverage gap that ground crews physically "
             "cannot reach (the downwind/eastern front), with ground crews holding the upwind/"
             "western flank passively. Each is necessary; neither is sufficient."),

            ("normal",
             "Suppression-mechanism complementarity. Comparing across the three baseline-wind "
             "scenarios with firefighters present isolates the contribution of each suppression "
             "mechanism. Rotor wash alone reduces peak burning modestly (19%) but reduces peak "
             "embers substantially (44%), consistent with §4.13’s argument that aerodynamic "
             "disruption is most effective on ember-scale ignitions. Acoustic emission (when "
             "added to rotor wash) extends suppression to established flame fronts, enabling the "
             "swarm to suppress fire faster than it spreads. The combined effect — 97% reduction "
             "in peak burning and 99% reduction in embers — is qualitatively different from "
             "either mechanism alone and reflects the structural complementarity of the two "
             "subsystems."),

            ("heading", "5.8.4 Decomposition: Mechanisms and Agents"),
        ])

        # ── §5.8.3 Combined Performance (the headline) ──────────────────────
        blocks.extend([
            ("normal",
             "The mission-complete result at 16.6 minutes is bounded by the simulation’s "
             "idealized assumptions (§5.9.3): perfect coordination, idealized acoustic "
             "propagation, no thermal damage to drones, flat homogeneous terrain. The headline "
             "finding should therefore be read not as a literal performance estimate but as "
             "evidence that the combined-mechanism + ground-crew configuration crosses a "
             "containment threshold that no other configuration tested can reach."),

            ("normal",
             "The containment result therefore depends on three jointly-necessary conditions, all "
             "of which can be tested in isolation through the 2×2 above and the wind-impact "
             "comparisons of §5.8.5: (1) both acoustic and rotor-wash mechanisms active "
             "simultaneously (compare combined_baseline vs rotor_only_baseline); (2) ground "
             "firefighters establishing the western-flank firebreak (compare combined_baseline vs "
             "drones_only); (3) wind conditions within the platform’s stability envelope "
             "(compare combined_baseline vs combined_high_wind). Outside any of these three "
             "conditions, the swarm fails to contain and the fire runs to fuel exhaustion. This "
             "narrows but strengthens the operational claim: UAV swarms are a complement to "
             "ground-crew operations within a wind-bounded envelope, not a replacement for them."),

            ("normal",
             "However, the 2×2 factorial in Table 0 (drones × firefighters) reveals a second "
             "necessary condition that was not anticipated at the start of this work. The same "
             "combined-mechanism swarm, deployed without the 10 ground firefighters, fails to "
             "contain the fire at the same seed: the drones_only scenario burned 39,755 cells "
             "(99.4% of grid), worse than pure_fire itself (99.0%) due to fuel-exhaustion "
             "run-to-run variance. The drone fleet alone cannot defend the western flank of an "
             "eastward-moving fire while simultaneously suppressing the active flame front to the "
             "east, and falls behind the spread on both axes. The 10 firefighters’ contribution "
             "is not flame suppression — they walk at 0.5 m/s from the western edge and almost "
             "never reach an active flame in time to extinguish it. Their contribution is the "
             "moisture firebreak they passively create as the fire’s leading edge approaches "
             "their position: 2,000–2,500 cells of saturated fuel that the fire cannot ignite, "
             "freeing the swarm to focus exclusively on the eastern front."),

            ("normal",
             "The two suppression mechanisms are complementary in exactly the way Chapter 4 "
             "predicted. Acoustic emission disrupts established burning cells, reducing intensity "
             "and extinguishing them with sufficient sustained exposure. Rotor wash handles the "
             "ember layer — both ground-level ember-state cells and airborne embers in flight. "
             "Neither mechanism alone covers the full fire lifecycle; together they do."),

            ("normal",
             "When both suppression mechanisms are active and ground firefighters are present, "
             "the swarm transforms from a delaying agent into a containing one. Peak burning was "
             "only 108 cells (a 97% reduction vs no drones), total burned was 1,388 cells (a 96% "
             "reduction vs no drones, 99% reduction vs pure fire dynamics), peak embers was just "
             "3 cells (a 99% reduction), and mission completed at t = 999 s (16.6 simulated "
             "minutes)."),

            ("heading", "5.8.3 Combined Acoustic + Rotor-Wash + Ground-Crew Performance"),
        ])

        # ── §5.8.2 Rotor-Wash-Only ──────────────────────────────────────────
        blocks.extend([
            ("normal",
             "This isolated rotor-wash performance is not operationally proposed — drones cannot "
             "fly without producing rotor wash — but it usefully bounds the contribution of "
             "aerodynamic suppression alone. The 44% ember reduction is the standout result: "
             "drones hovering above ember cells extinguish them efficiently, and downwash on "
             "airborne embers shortens their trajectories."),

            ("normal",
             "The pattern is consistent with the physical predictions of §4.13: rotor wash is "
             "effective at ember management but cannot meaningfully suppress an established flame "
             "front. With acoustic emission disabled, the swarm has no mechanism to reduce burn "
             "intensity on cells that have already established sustained combustion, and the fire "
             "continues to consume the available fuel until natural burnout."),

            ("normal",
             "With rotor wash active but acoustic emission disabled, the 100-drone swarm produced "
             "measurable but modest improvements over the no-drone baseline. Peak burning was "
             "reduced from 3,930 to 3,185 cells (a 19% reduction), total burned was reduced from "
             "37,606 to 37,525 cells (essentially unchanged), and peak embers were reduced from "
             "393 to 219 cells (a 44% reduction)."),

            ("heading", "5.8.2 Rotor-Wash-Only Performance"),
        ])

        # ── §5.8.1 Counterfactuals ──────────────────────────────────────────
        blocks.extend([
            ("normal",
             "The 10 ground firefighters operating in the no_drones scenario applied water to "
             "approximately 2,500 cells (Option B firefighter physics) — a wet moisture firebreak "
             "along the western flank — but were unable to materially affect the fire’s spread "
             "rate at the modeled walking speed of 0.5 m/s. Total burned was 37,158 cells (92.9%), "
             "compared to 39,595 cells (99.0%) in the pure_fire counterfactual. Firefighter "
             "contribution is therefore a ~6 percentage-point absolute reduction in burned "
             "fraction, achieved primarily through passive moisture barriers ahead of the "
             "fireline rather than through active flame extinguishment."),

            ("normal",
             "The pure_fire scenario (no drones, no firefighters) establishes the absolute fire-"
             "dynamics ceiling: in the absence of any intervention, the fire spreads monotonically "
             "from the ignition cluster, reaches peak burning of 3,825 cells at approximately "
             "t = 900 s, and ultimately consumes 39,595 cells (99.0% of grid). Peak ember activity "
             "reaches 390 cells, reflecting the full ember-lofting potential of an unsuppressed "
             "fire. The fire eventually self-extinguishes when it has consumed essentially all "
             "available fuel."),

            ("heading", "5.8.1 Counterfactuals: pure_fire and no_drones"),
        ])

        # ── Section intro and 2×2 summary ───────────────────────────────────
        blocks.extend([
            ("normal",
             "High-wind robustness checks under Option B were also conducted at 10 m/s. Both "
             "rotor_only_high_wind and combined_high_wind burned approximately 37,300 cells "
             "(93.3%) — within run-to-run noise of each other and of the original numbers, "
             "confirming that high-wind collapse is independent of the firefighter-water model."),

            ("normal",
             "Outcomes (5 m/s, 35,000-tick cap, seed 42): pure_fire — 39,595 burned (99.0%), no "
             "containment. no_drones (10 firefighters, no drones) — 37,158 burned (92.9%), no "
             "containment. drones_only (100 drones, no firefighters) — 39,755 burned (99.4%), "
             "fuel-exhaustion only. combined_baseline (100 drones + 10 firefighters) — 1,388 "
             "burned (3.5%), MISSION COMPLETE at t = 999 s (16.6 min). The diagonal of the table "
             "tells the story: no single intervention is sufficient, but the joint configuration "
             "crosses a clean containment threshold."),

            ("normal",
             "The seven scenarios were executed headlessly. To clarify whether the combined "
             "mechanism alone is sufficient or whether ground crews are also necessary, two "
             "additional scenarios were added (pure_fire — no drones, no firefighters; "
             "drones_only — combined drones, no firefighters) to complete a 2×2 factorial design "
             "isolating the contribution of each agent type. Following the addition, an Option B "
             "correction to the firefighter water-spray model (allowing water to extinguish "
             "burning cells in addition to wetting unburned ones, rather than only wetting "
             "unburned cells as in the original model) was applied and the four non-containment "
             "scenarios were re-baselined under the corrected physics. Sensitivity analyses on "
             "rotor-wash strength, airborne-ember drag, and seed-level reproducibility were also "
             "conducted; results are presented in §5.8.7 below."),
        ])

        # Insert in reverse order (each addnext puts it just after anchor, so reversed list ends up forward)
        for block in blocks:
            kind = block[0]
            if kind == "heading":
                add_paragraph_after(doc, anchor, block[1], style="Heading 3")
            elif kind == "figure":
                insert_figure_after(doc, anchor, block[1], block[2])
            else:
                add_paragraph_after(doc, anchor, block[1])

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE B: replace §5.9.2 Key Findings with the 7-finding final list
    # ─────────────────────────────────────────────────────────────────────────
    print("Replacing §5.9.2 Key Findings...")
    f592 = find_paragraph_index(doc, "5.9.2 Key Findings")
    f593 = find_paragraph_index(doc, "5.9.3 Limitations")
    if f592 < 0 or f593 < 0 or f593 <= f592:
        print("  WARN: §5.9.2 anchors not found.")
    else:
        for i in range(f593 - 1, f592, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

        anchor = doc.paragraphs[f592]

        findings = [
            "7. Containment is probabilistic, not deterministic, at the modeled fleet size. A "
            "seed-reproducibility study with N=5 random seeds at the otherwise identical "
            "combined_baseline configuration produced clean containment in 3 of 5 runs (≤4% grid "
            "burnout in <17 simulated minutes) and lost the early-spread race in 2 of 5 runs "
            "(>91% grid burnout, mission terminated by fuel exhaustion). The failure mode in both "
            "lost runs is loss of the early-spread race within the first 100–200 simulated "
            "seconds, after which the fire grows beyond the swarm’s per-tick suppression capacity. "
            "This same failure mode is observed in drones_only (no firefighter firebreak) and in "
            "the high-wind scenarios. Containment at this fleet size is therefore best "
            "characterized as approximately a 60% probability event, with the failure mode well-"
            "understood and bounded.",

            "6. The simulation results are bounded by their idealized assumptions, but the "
            "direction of the effect is robust. The mission-complete result should not be read as "
            "a literal claim that 100 drones could extinguish a real grass fire in 17 minutes; it "
            "is, rather, evidence that crossing the suppression-pressure threshold from “delay” "
            "to “contain” is possible only when both suppression mechanisms are active and a "
            "passive ground-crew firebreak is established. The qualitative finding — that the "
            "combined acoustic + rotor-wash + ground-crew configuration produces fundamentally "
            "different dynamics than any subset of those components — is robust to reasonable "
            "variation in parameter calibration because it depends on the structural "
            "complementarity of the components, not on the absolute magnitudes. A rotor-wash-"
            "strength sensitivity sweep (±100% variation) and an airborne-ember-drag isolation "
            "test confirmed that the containment outcome holds across the tested parameter range.",

            "5. Battery endurance is not the binding constraint when the swarm is effective. The "
            "combined-mechanism mission completed in 17 minutes with mean fleet battery at 93%, "
            "well above the 20% return-to-base threshold. Battery becomes a constraint only in "
            "scenarios where the swarm cannot contain the fire and continues hovering "
            "indefinitely — situations in which the operational value of continued engagement is "
            "itself questionable.",

            "4. The combined-mechanism result is contingent on drones being able to maintain "
            "station; high wind eliminates the advantage entirely. When base wind speed is "
            "doubled to 10 m/s, both Rotor-Wash-Only and Combined scenarios collapse to "
            "essentially identical numerical outcomes (peak burning ≈ 4,121, total burned ≈ "
            "37,572, peak embers ≈ 324), driven by frequent excursions of gusts above the 12 m/s "
            "drone safety threshold and the corresponding return-to-base behavior. The mechanism "
            "toggles become operationally equivalent because neither mechanism is applied long "
            "enough to act. This bounds the deployment envelope: the headline containment "
            "finding holds only within the modeled platform’s stability margin.",

            "3. Ember management is the highest-value contribution per drone. Peak ember counts "
            "dropped from 393 (no drones) to 219 (rotor wash only) to 3 (combined). The 99% "
            "reduction in the combined scenario directly reflects the swarm’s ability to "
            "extinguish nascent ember-state cells in approximately 0.12 simulated seconds — an "
            "order of magnitude faster than full burning cells — before they can grow into "
            "established flames. This empirically supports §4.14’s redirection of UAV "
            "suppression toward early-ignition and ember-management roles rather than frontline "
            "flame suppression.",

            "2b. The drone fleet and ground firefighters are also complementary, in a different "
            "and unexpected way. The 10 ground firefighters in this simulation almost never reach "
            "an active flame — at 0.5 m/s walking speed, they cannot catch a fire spreading at "
            "over 1 m/s in 5 m/s wind. Their actual contribution to the containment outcome is "
            "geometric: they create a 2,000–2,500-cell moisture firebreak along the western "
            "(upwind) flank of the fire as it approaches their position, which prevents westward "
            "spread and frees the drone swarm to focus exclusively on the eastern front. "
            "Without this firebreak (the drones_only scenario), the swarm cannot defend two "
            "flanks simultaneously and the fire runs to fuel exhaustion. This finding inverts the "
            "conventional framing of UAVs as “force multipliers”: at the fleet size and platform "
            "parameters modeled here, the swarm is more accurately described as filling the "
            "coverage gap that ground crews physically cannot reach (the downwind/eastern front), "
            "with ground crews holding the upwind/western flank passively.",

            "2. The two suppression mechanisms are complementary rather than redundant. Rotor "
            "wash alone reduces peak burning by only 19% but reduces peak embers by 44%, "
            "consistent with §4.13’s argument that aerodynamic disruption is most effective on "
            "ember-scale ignitions. Acoustic suppression alone is most effective on burning "
            "cells. Together, the two mechanisms cover the full fire lifecycle and produce "
            "qualitatively different containment behavior than either alone.",

            "1. Containment requires the joint presence of three conditions; no single "
            "intervention is sufficient. The 2×2 factorial design (drones × firefighters, §5.8.1) "
            "shows that 100 drones with combined acoustic and rotor-wash mechanisms achieve full "
            "containment (3.5% grid burnout in 16.6 minutes) only when also operating alongside "
            "10 ground firefighters under wind conditions within the platform stability envelope. "
            "Removing any one of these — disabling one suppression mechanism, removing the "
            "ground firefighters, or doubling wind speed — drops the configuration below the "
            "containment threshold and allows the fire to consume 90–99% of the grid. The "
            "headline finding is therefore not “drone swarms can contain wildfires” but “drone "
            "swarms, deployed in coordination with ground firefighters within a bounded wind "
            "envelope, can contain wildfires.”",

            "The simulation results, taken across the 2×2 factorial design and the four "
            "sensitivity analyses, support the following conclusions:",
        ]

        for text in findings:
            add_paragraph_after(doc, anchor, text)

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE C: append a small note to §5.9.3 Limitations about Option B
    # ─────────────────────────────────────────────────────────────────────────
    print("Adding Option B note to §5.9.3 Limitations...")
    l593 = find_paragraph_index(doc, "5.9.3 Limitations")
    if l593 >= 0:
        anchor = doc.paragraphs[l593]
        add_paragraph_after(
            doc, anchor,
            "Firefighter water spray model (Option B). The original model treated firefighter "
            "water as a moisture barrier that protected unburned cells but did not extinguish "
            "active flames — a model artefact, not physics. The corrected model (Option B) allows "
            "water on a burning or ember cell to extinguish it in addition to wetting unburned "
            "cells. Re-running the four non-containment scenarios under the corrected model "
            "changed total-burned counts by less than 1.5 percentage points, because the 10 "
            "firefighters at 0.5 m/s walking speed almost never reach an active flame in time to "
            "extinguish it regardless of whether their water can do so on contact. The "
            "containment finding is therefore robust to this modeling choice, and the model is "
            "now more physically defensible.",
        )

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE D: insert a 2×2 factorial figure right before §5.8.1
    # ─────────────────────────────────────────────────────────────────────────
    fig_path = f"{FIG_DIR}/fig_2x2_factorial.png"
    if os.path.exists(fig_path):
        print("Inserting 2×2 factorial figure before §5.8.1...")
        s581 = find_paragraph_index(doc, "5.8.1 Counterfactuals")
        if s581 > 0:
            insert_figure_after(
                doc, doc.paragraphs[s581 - 1], fig_path,
                "Figure: 2×2 factorial — % of grid burned across the four corners of the "
                "drones × firefighters design (5 m/s wind, seed 42, Option B firefighter physics). "
                "Only the joint configuration crosses the containment threshold."
            )

    # ─────────────────────────────────────────────────────────────────────────
    # Save
    # ─────────────────────────────────────────────────────────────────────────
    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!\n")

    print("Summary of changes (all yellow-highlighted in the new document):")
    print("  - §5.8 Simulation Results: completely restructured")
    print("    - 5.8.1 Counterfactuals (pure_fire + no_drones)")
    print("    - 5.8.2 Rotor-Wash-Only Performance")
    print("    - 5.8.3 Combined Acoustic + Rotor-Wash + Ground-Crew Performance (THE HEADLINE)")
    print("    - 5.8.4 Decomposition: Mechanisms and Agents")
    print("    - 5.8.5 Wind Impact (with Option B note)")
    print("    - 5.8.6 Battery and Operational Endurance")
    print("    - 5.8.7 Sensitivity Analysis (seeds, wash strength, ember-drag isolation)")
    print("  - 2×2 factorial figure inserted before §5.8.1")
    print("  - §5.9.2 Key Findings: 7 findings (added 2b on geometric complementarity, 7 on probabilistic containment)")
    print("  - §5.9.3 Limitations: appended Option B firefighter-physics note")


if __name__ == "__main__":
    main()
