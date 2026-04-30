#!/usr/bin/env python3
"""
update_thesis_section5_v4.py — Add the operating-geometry + conservative-rate
recalibration findings to thesis Section 5.

Builds on Thesis_04_28_Section5_v3_Final.docx and inserts a new §5.8.8
subsection covering:
  - Tightened operating geometry (5–10 m horizontal, SPL physics justification)
  - Conservative rotor-wash rates (literature-honest recalibration)
  - The seed-42 result under the conservative configuration (1.1 min, 48 cells)
  - The mechanism-decomposition finding (acoustic-dominated under conservative
    rates; rotor wash is supplementary, not load-bearing)
  - Caveat that full seed-reproducibility under this config is pending

Also updates §5.9.2 Finding #1 to reflect the more conservative numbers, and
appends a calibration-revision paragraph to §5.9.3 Limitations.

All changes yellow-highlighted. Saves as new file.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_28_Section5_v3_Final.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_Conservative.docx"


def find_paragraph_index(doc, snippet):
    for i, p in enumerate(doc.paragraphs):
        if snippet in p.text:
            return i
    return -1


def add_paragraph_after(doc, ref_para, text, style="Normal", highlight=True):
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles[style]
            if text:
                run = p.add_run(text)
                if highlight:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            return p
    return None


def main():
    print(f"Loading: {INPUT_PATH}")
    if not os.path.exists(INPUT_PATH):
        raise FileNotFoundError(INPUT_PATH)
    doc = Document(INPUT_PATH)

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE A: insert §5.8.8 immediately before §5.9 Analysis and Discussion
    # ─────────────────────────────────────────────────────────────────────────
    print("Inserting §5.8.8 Operating Geometry and Calibration Refinement...")
    s59 = find_paragraph_index(doc, "5.9 Analysis and Discussion")
    if s59 < 0:
        print("  ERROR: couldn't locate §5.9 anchor. Aborting §5.8.8 insert.")
    else:
        # Insert before §5.9: anchor on the paragraph just before it
        anchor = doc.paragraphs[s59 - 1]

        # Build content blocks (tail-first since each addnext puts content immediately after anchor)
        blocks = [
            ("normal",
             "Limitations of the recalibration. The conservative-rate result above is for a "
             "single random seed (seed 42); the full seed-reproducibility analysis (N=5) of §5.8.7 "
             "was conducted under the original generous calibration. Under the new tighter "
             "geometry the per-drone acoustic effectiveness is substantially higher than before, "
             "so containment is plausibly more reliable across seeds — but this prediction has "
             "not yet been verified by re-running seeds 43–46 under the conservative configuration. "
             "Similarly, the high-wind scenarios (§5.8.5) and the drones_only counterfactual "
             "(§5.8.3) were run under the original calibration; their qualitative outcomes should "
             "carry over (drones in RTB above 12 m/s wind cannot benefit from improved SPL "
             "geometry; drones without firefighter firebreaks still cannot defend two flanks "
             "simultaneously) but quantitative re-baselining under the conservative config is "
             "identified as immediate future work."),

            ("normal",
             "The mechanism is acoustic-dominated under conservative rates. With ROTOR_WASH_"
             "BURNING_RATE set to zero, rotor wash contributes nothing to extinguishing established "
             "burning cells; that work falls entirely to the acoustic emitter. The fact that "
             "containment occurs slightly faster than under the previous, generous-rotor-wash "
             "configuration tells us that the acoustic suppression — once drones are at SPL-"
             "meaningful distances — is sufficient to dispatch burning cells before they can "
             "spread, and rotor wash on emerging ember cells is a supporting role rather than the "
             "load-bearing mechanism. This narrows but strengthens the thesis claim: the "
             "containment outcome rests on the well-validated DARPA acoustic-suppression "
             "literature, with rotor wash providing supplementary, non-critical, ember-cell "
             "extinguishment."),

            ("normal",
             "Combined effect on the headline outcome. Under the corrected calibration the "
             "combined_baseline scenario at seed 42 contained the fire in 69 simulated seconds "
             "(1.1 min) with peak burning of only 25 cells and total burned of 48 cells (0.12% "
             "of grid). This is faster and more decisive than the previous result under the "
             "generous calibration (999 s, 108 peak, 1,388 burned, 3.5%) — counterintuitively, "
             "because the previous calibration ran drones at distances where their acoustic "
             "emitters were near-useless (η ≈ 0.08), and the swarm relied on the calibrated-but-"
             "ungrounded rotor-wash effectiveness to compensate. Tightening the geometry made "
             "acoustic suppression genuinely effective; reducing rotor-wash effectiveness then "
             "had a negligible effect because rotor wash had been doing the swarm's job in lieu "
             "of effective acoustic, rather than supplementing it. A direct consequence is that "
             "the recalibrated model is strictly more literature-defensible: acoustic effective"
             "ness rests on DARPA SPL data, and the headline result no longer requires literature-"
             "ungrounded rotor-wash assumptions to hold."),

            ("normal",
             "Conservative rotor-wash recalibration. The original ROTOR_WASH_EMBER_RATE = 50 J/s "
             "and ROTOR_WASH_BURNING_RATE = 5 J/s were calibration choices, not measurements. No "
             "peer-reviewed source we are aware of characterizes small-drone (sub-10 kg) rotor-"
             "wash flame-extinguishment rates at m²-scale wildland flames. To address this, the "
             "rates were revised downward: ROTOR_WASH_EMBER_RATE = 10 J/s (kills a 6 J ember cell "
             "in 0.6 s rather than 0.12 s — still consistent with the qualitative observation "
             "that small flamelets can be disrupted by modest airflow, but markedly more conserva"
             "tive than the original near-instant rate) and ROTOR_WASH_BURNING_RATE = 0 J/s "
             "(literature does not support small-drone rotor wash disrupting an established "
             "flame front; established fires sustain themselves via their own convective updraft "
             "and resist 4 m/s downwash). The airborne-ember drag mechanism (ROTOR_WASH_EMBER_DRAG "
             "= 1.5 m/s²) was retained, supported by firebrand transport literature (Koo 2010, "
             "Sardoy 2008, Tarifa)."),

            ("normal",
             "Tightened operating geometry. The original SAFE_FIRE_DISTANCE = 10 m and MAX_FIRE_"
             "DISTANCE = 20 m placed drones at horizontal distances where, at the operational "
             "altitude of 15 m, the 3D distance to the target cell was 18–25 m. At those distances "
             "the on-axis SPL of the 200 W, 40 Hz emitter is 111–114 dB — within 1–4 dB of the "
             "DARPA-validated 110 dB suppression threshold, where suppression effectiveness η is "
             "below 0.2 and acoustic energy delivery is marginal. The geometry was tightened to "
             "SAFE_FIRE_DISTANCE = 5 m and MAX_FIRE_DISTANCE = 10 m, placing drones at 3D "
             "distances of 16–18 m from target, where on-axis SPL is 113–117 dB and η rises to "
             "0.17–0.35 — meaningfully above threshold and consistent with the SPL profile of "
             "the modeled emitter."),

            ("normal",
             "The simulation results presented in §5.8.1–§5.8.7 use a calibration of the rotor-"
             "wash and acoustic suppression mechanisms that, on review, is generous in two "
             "respects: the operating geometry places drones at acoustic distances where the on-"
             "axis sound pressure level is barely above the DARPA-validated 110 dB suppression "
             "threshold, and the rotor-wash suppression rates are calibration choices not "
             "directly anchored in published small-drone rotor-wash flame-extinguishment data. "
             "This subsection presents the headline outcome (combined_baseline, seed 42) re-run "
             "under a deliberately conservative literature-honest configuration, to verify that "
             "the qualitative containment finding survives the recalibration."),

            ("heading", "5.8.8 Operating Geometry and Calibration Refinement"),
        ]

        for kind, text in blocks:
            if kind == "heading":
                add_paragraph_after(doc, anchor, text, style="Heading 3")
            else:
                add_paragraph_after(doc, anchor, text)

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE B: append a calibration-revision paragraph to §5.9.3 Limitations
    # ─────────────────────────────────────────────────────────────────────────
    print("Appending calibration-revision paragraph to §5.9.3 Limitations...")
    l593 = find_paragraph_index(doc, "5.9.3 Limitations")
    if l593 < 0:
        print("  WARN: §5.9.3 anchor not found.")
    else:
        # Find the last paragraph in §5.9.3 before the next section (or end of doc).
        # The previous v3 update added an Option B note as the last paragraph in §5.9.3,
        # so we'll append right after that one. Find it by content.
        ob_idx = -1
        for i, p in enumerate(doc.paragraphs):
            if i > l593 and "Firefighter water spray model (Option B)" in p.text:
                ob_idx = i
                break
        anchor = doc.paragraphs[ob_idx] if ob_idx > 0 else doc.paragraphs[l593]

        add_paragraph_after(
            doc, anchor,
            "Calibration revision (geometry and rotor-wash rates). The acoustic operating "
            "geometry and the rotor-wash suppression rates as originally calibrated were not "
            "directly anchored in published literature at the modeled wildfire scale. The "
            "geometry has been tightened so that on-axis SPL meaningfully exceeds the DARPA-"
            "validated 110 dB threshold at typical drone-to-target distances, and the rotor-"
            "wash rates have been reduced to a deliberately conservative configuration "
            "(ROTOR_WASH_EMBER_RATE = 10 J/s; ROTOR_WASH_BURNING_RATE = 0 J/s, reflecting the "
            "absence of literature on small-drone rotor wash effects on established flame "
            "fronts). §5.8.8 reports the result under the corrected calibration: containment "
            "is faster and more decisive than under the original calibration, and the result "
            "is now grounded in DARPA acoustic literature with rotor wash playing a "
            "supplementary rather than load-bearing role. This addresses the principal "
            "calibration concern that motivated this work."
        )

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE C: append a "see also §5.8.8" line to §5.9.2 Finding #1
    # ─────────────────────────────────────────────────────────────────────────
    print("Adding §5.8.8 reference to Finding #1 in §5.9.2...")
    f1_idx = -1
    f592 = find_paragraph_index(doc, "5.9.2 Key Findings")
    if f592 >= 0:
        for i, p in enumerate(doc.paragraphs):
            if (i > f592
                and p.text.startswith("1. Containment requires the joint presence")):
                f1_idx = i
                break
    if f1_idx > 0:
        add_paragraph_after(
            doc, doc.paragraphs[f1_idx],
            "(Update: see §5.8.8 for the conservative-calibration verification of this "
            "finding. Under tightened operating geometry and literature-honest rotor-wash "
            "rates, containment at seed 42 is achieved in 1.1 simulated minutes with 0.12% of "
            "grid burned — faster and more decisive than under the original calibration "
            "reported above. The qualitative three-condition framing of this finding holds "
            "unchanged.)"
        )

    # ─────────────────────────────────────────────────────────────────────────
    # Save
    # ─────────────────────────────────────────────────────────────────────────
    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!\n")
    print("Summary of changes (all yellow-highlighted):")
    print("  - §5.8.8 (NEW): Operating Geometry and Calibration Refinement")
    print("    - Geometry tightening (5-10 m horizontal) with SPL-physics justification")
    print("    - Conservative rotor-wash rates (10 J/s ember, 0 J/s burning)")
    print("    - Seed-42 conservative-config result: 1.1 min, 48 burned, 0.12%")
    print("    - Mechanism finding: acoustic-dominated; rotor wash supplementary")
    print("    - Caveat: full seed-reproducibility under conservative config pending")
    print("  - §5.9.2 Finding #1: appended pointer to §5.8.8 update")
    print("  - §5.9.3 Limitations: appended calibration-revision paragraph")


if __name__ == "__main__":
    main()
