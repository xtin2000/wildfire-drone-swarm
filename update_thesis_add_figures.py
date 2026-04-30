#!/usr/bin/env python3
"""Insert three figures into the thesis:
  - fig_timeline_overlay.png and fig_fleet_state.png inside §5.8.1
  - fig_high_wind_collapse.png inside §5.8.5 (the Wind Impact subsection;
    note: §5.8.4 in the current thesis numbering is "Decomposition", not
    Wind Impact — so this figure lands in §5.8.5 where it textually belongs)

Builds on Thesis_04_29_Section5_v5.docx. Saves as new file.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_v5.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_v6.docx"
FIG_DIR = "/Users/xtin2000/wildfire_drone_swarm/results_v2/figures"

CAPTION_TIMELINE = (
    "Figure: Burning-cell trajectory over simulated time across all five major scenarios "
    "(5 m/s wind, seed 42). The pure_fire and rotor_only_baseline trajectories climb to peaks "
    "above 3,000 cells and run for the full 35,000-tick window. The drones_only trajectory "
    "(combined mechanisms but no firefighter firebreak) also fails to contain. Only the "
    "combined_baseline configuration (drones + firefighters) suppresses burning to near zero "
    "within the first 17 simulated minutes and exits via mission_complete."
)

CAPTION_FLEET = (
    "Figure: Drone fleet state distribution over simulated time for combined_baseline "
    "(seed 42, 5 m/s wind). Top panel shows the count of drones in each state (Idle / Transit / "
    "Hovering / Emitting / Returning / Charging) across the run; the swarm transitions rapidly "
    "from Idle into Transit and Emitting as the fire is detected, with sustained Emitting at the "
    "fire perimeter until containment. Bottom panel shows mean and minimum fleet battery and "
    "mean platform stability over the same window — battery drains less than 7% over the "
    "16.6-minute mission and stability remains comfortably above the 12 m/s safety threshold."
)

CAPTION_WIND = (
    "Figure: Cumulative cells burned over simulated time, comparing rotor-only and combined "
    "configurations at 5 m/s (solid) versus 10 m/s (dashed) wind. At 5 m/s the combined "
    "mechanism (green solid) achieves containment near 1,000 simulated seconds while the "
    "rotor-only configuration (blue solid) runs to fuel exhaustion. At 10 m/s the two "
    "configurations collapse onto each other (green dashed and blue dashed are visually "
    "indistinguishable) — both scenarios converge to ~93% grid burnout because gusts above "
    "the 12 m/s safety threshold force frequent return-to-base behaviour, eliminating the "
    "mechanism advantage."
)


def find_paragraph_index(doc, snippet):
    for i, p in enumerate(doc.paragraphs):
        if snippet in p.text:
            return i
    return -1


def insert_figure_after(doc, ref_para, fig_path: str, caption: str) -> None:
    """Insert a centered figure followed by a centered yellow-highlighted italic caption."""
    # We use addnext, which inserts immediately after the ref. To place [figure][caption]
    # in correct order with the figure first, we addnext the caption first, then addnext
    # the figure (which will land between the ref and the caption).

    cap_p = OxmlElement("w:p")
    ref_para._element.addnext(cap_p)
    for p in doc.paragraphs:
        if p._element is cap_p:
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(10)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            break

    if os.path.exists(fig_path):
        fig_p = OxmlElement("w:p")
        ref_para._element.addnext(fig_p)
        for p in doc.paragraphs:
            if p._element is fig_p:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(fig_path, width=Inches(5.5))
                break
    else:
        print(f"WARN: figure not found: {fig_path}")


def main() -> None:
    print(f"Loading: {INPUT_PATH}")
    if not os.path.exists(INPUT_PATH):
        raise FileNotFoundError(INPUT_PATH)
    doc = Document(INPUT_PATH)

    # ── Insert fleet-state and timeline figures into §5.8.1 ──
    # Find the §5.8.1 heading and the §5.8.2 heading; insert just before §5.8.2
    s581 = find_paragraph_index(doc, "5.8.1 Counterfactuals")
    s582 = find_paragraph_index(doc, "5.8.2 Rotor-Wash-Only Performance")
    if s581 < 0 or s582 < 0:
        raise RuntimeError("Could not locate §5.8.1 / §5.8.2 anchors")
    print(f"§5.8.1 at {s581}, §5.8.2 at {s582}")

    # Anchor: paragraph immediately before §5.8.2 heading (last paragraph of §5.8.1)
    anchor_581 = doc.paragraphs[s582 - 1]
    # Insert order: timeline first (will appear first), then fleet-state below it.
    # Because each insert_figure_after pushes its content immediately after the anchor,
    # we insert fleet-state FIRST so it ends up below timeline.
    insert_figure_after(doc, anchor_581, f"{FIG_DIR}/fig_fleet_state.png", CAPTION_FLEET)
    insert_figure_after(doc, anchor_581, f"{FIG_DIR}/fig_timeline_overlay.png", CAPTION_TIMELINE)

    # ── Insert wind-impact figure into §5.8.5 (Wind Impact) ──
    s585 = find_paragraph_index(doc, "5.8.5 Wind Impact")
    s586 = find_paragraph_index(doc, "5.8.6 Battery and Operational Endurance")
    if s585 < 0 or s586 < 0:
        raise RuntimeError("Could not locate §5.8.5 / §5.8.6 anchors")
    print(f"§5.8.5 at {s585}, §5.8.6 at {s586}")

    anchor_585 = doc.paragraphs[s586 - 1]
    insert_figure_after(doc, anchor_585, f"{FIG_DIR}/fig_high_wind_collapse.png", CAPTION_WIND)

    # Save
    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!\n")
    print("Summary:")
    print("  - §5.8.1 Counterfactuals: inserted fig_timeline_overlay.png + fig_fleet_state.png")
    print("    with yellow-highlighted captions")
    print("  - §5.8.5 Wind Impact: inserted fig_high_wind_collapse.png with caption")
    print("    NOTE: §5.8.4 in the current thesis is 'Decomposition: Mechanisms and Agents';")
    print("    the wind impact subsection is §5.8.5. Figure was placed where it textually")
    print("    belongs. If you want it in §5.8.4 instead, let me know and I'll move it.")


if __name__ == "__main__":
    main()
