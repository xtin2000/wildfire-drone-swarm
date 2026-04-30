#!/usr/bin/env python3
"""
update_thesis_long_run.py — Insert long-run findings into thesis Section 5.

Updates three specific places with new battery and time-to-damage data,
highlighting all changes in yellow so the user can review them in Word.

Creates a NEW file (does not overwrite the original).
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_20_ChristineLangmayr.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_26_Long_Run_Updates.docx"
NEW_BATTERY_FIG = "/Users/xtin2000/wildfire_drone_swarm/results_long/figures/fig_battery.png"


def add_highlighted_paragraph(doc, after_para, text, style="Normal"):
    """Insert a new paragraph after the given one, with all text yellow-highlighted."""
    new_p_element = OxmlElement("w:p")
    after_para._element.addnext(new_p_element)
    new_para = None
    for p in doc.paragraphs:
        if p._element is new_p_element:
            new_para = p
            break
    if new_para is None:
        return None
    new_para.style = doc.styles[style]
    if text:
        run = new_para.add_run(text)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return new_para


def insert_figure(doc, after_para, fig_path, caption):
    """Insert a centered figure with caption after the given paragraph."""
    cap_p = OxmlElement("w:p")
    after_para._element.addnext(cap_p)
    for p in doc.paragraphs:
        if p._element is cap_p:
            p.style = doc.styles["Normal"]
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(10)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    if os.path.exists(fig_path):
        fig_p = OxmlElement("w:p")
        after_para._element.addnext(fig_p)
        for p in doc.paragraphs:
            if p._element is fig_p:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(fig_path, width=Inches(5.5))
                break


def replace_paragraph_text(para, new_text, highlight=True):
    """Clear a paragraph and replace its text, optionally highlighted."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
        if highlight:
            para.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        run = para.add_run(new_text)
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def find_paragraph_by_text(doc, snippet):
    """Return the first paragraph whose text contains the given snippet."""
    for p in doc.paragraphs:
        if snippet in p.text:
            return p
    return None


def main():
    print(f"Loading: {INPUT_PATH}")
    doc = Document(INPUT_PATH)

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE 1: §5.8.5 Battery and Operational Endurance
    # Replace the existing two paragraphs with the long-run analysis.
    # ─────────────────────────────────────────────────────────────────────────
    print("Updating §5.8.5 (battery analysis)...")

    para1 = find_paragraph_by_text(
        doc, "Battery depletion follows a characteristic pattern"
    )
    para2 = find_paragraph_by_text(
        doc, "The minimum battery curves show that the earliest"
    )

    if para1 is None or para2 is None:
        print("  ERROR: Could not find §5.8.5 paragraphs")
        return

    new_text_1 = (
        "Within the standard 5,000-tick (500-second) simulation window, battery levels remained "
        "above 90% across all scenarios because fleet engagement begins approximately 25 seconds "
        "after ignition. To characterize battery dynamics over operationally meaningful durations, "
        "the baseline and high-wind scenarios were also run for an extended 35,000 ticks "
        "(3,500 simulated seconds, or roughly 58 minutes of mission time)."
    )

    new_text_2 = (
        "Under the extended baseline run, mean fleet battery decreased approximately linearly "
        "from 100% to 61% over the 58-minute simulated mission, with the minimum-charged drone "
        "at 60.6%. The depletion rate of approximately 0.67% per simulated minute reflects "
        "sustained acoustic emission across the fleet, where the additional 250 W per drone is "
        "added to the 360 W base draw. No return-to-base events occurred during this window: "
        "the first RTB events are projected to begin around the 80–90 minute mark when the "
        "earliest-deployed drones reach the 20% charge threshold."
    )

    replace_paragraph_text(para1, new_text_1, highlight=True)
    replace_paragraph_text(para2, new_text_2, highlight=True)

    # Insert two additional paragraphs after para2 (in reverse order so they appear correctly)
    additional_para_2 = (
        "These extended results refine the original finding: for short-duration engagements at "
        "the 400 m × 400 m scale, battery endurance is not the binding constraint. For sustained "
        "operations beyond approximately 90 minutes, however, RTB cycling and the resulting "
        "coverage gaps would become a primary determinant of fleet effectiveness, and the "
        "1,500 Wh capacity should be considered a hard lower bound for any operational deployment."
    )

    additional_para_1 = (
        "In the extended high-wind scenario, mean battery declined far less, finishing at 92.8% "
        "after the same 58 minutes. This counterintuitive result is not a feature but a failure "
        "indicator: under 10 m/s wind, drones spent the majority of the mission in TRANSIT and "
        "RETURNING states with stability factors below 0.1, unable to enter the high-power "
        "EMITTING state for sustained periods. The high-wind scenario therefore demonstrates "
        "that wind constraints disable acoustic suppression long before battery becomes the "
        "binding factor."
    )

    add_highlighted_paragraph(doc, para2, additional_para_2)
    add_highlighted_paragraph(doc, para2, additional_para_1)

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE 2: Update the battery figure caption to note extended runs
    # The original Figure 6 caption stays; we ADD a note that an updated
    # version is available in §5.8.5
    # ─────────────────────────────────────────────────────────────────────────
    print("Adding extended-run battery figure...")

    # Find a place to insert the new figure — right after the new paragraphs
    # we just added. We'll do this by finding the paragraph matching our new text.
    insertion_anchor = None
    for p in doc.paragraphs:
        if p.text.startswith("These extended results refine"):
            insertion_anchor = p
            break

    if insertion_anchor:
        insert_figure(
            doc,
            insertion_anchor,
            NEW_BATTERY_FIG,
            "Figure 6 (Updated): Battery depletion over the extended 35,000-tick run "
            "(3,500 simulated seconds). Baseline depletes linearly to 61%; high-wind "
            "barely depletes because drones cannot sustain the EMITTING state.",
        )

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE 3: §5.9.2 Finding #5 (battery is not binding constraint)
    # ─────────────────────────────────────────────────────────────────────────
    print("Updating §5.9.2 finding #5...")

    finding_5 = find_paragraph_by_text(
        doc, "Battery endurance is not the binding constraint at this scale"
    )
    if finding_5:
        new_finding_5 = (
            "5. Battery endurance becomes binding only beyond approximately 90 minutes of "
            "sustained operation. The 1,500 Wh capacity provided ample endurance throughout "
            "both the 500-second baseline and the 3,500-second extended runs (final mean "
            "battery 61% under continuous emission). The primary short-duration limitations "
            "are acoustic range and fleet coordination efficiency; battery cycling and "
            "coverage gaps become the dominant constraint for missions extending beyond "
            "approximately 1.5 hours."
        )
        replace_paragraph_text(finding_5, new_finding_5, highlight=True)
    else:
        print("  WARNING: Could not find finding #5 to update")

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE 4: Add new finding #7 (time-to-equivalent-damage)
    # ─────────────────────────────────────────────────────────────────────────
    print("Adding new finding #7 (time-to-equivalent-damage)...")

    # Find finding #6 (last existing finding) to insert after
    finding_6 = find_paragraph_by_text(
        doc, "Ember control is the swarm"
    )
    if finding_6:
        new_finding_7 = (
            "7. Drones delay rather than contain wildfire spread. Extended 35,000-tick runs "
            "reveal that under sustained 100-drone engagement, the fire eventually consumes "
            "the same fraction of the grid as the uncontrolled scenario (37,171 vs 37,606 "
            "cells, a 1.2% difference at t=3,500 s). The operational value of the swarm is "
            "therefore measured in time bought rather than damage prevented: drones delayed "
            "equivalent fire damage by approximately 35 minutes (2,131 simulated seconds). "
            "This delay is the operationally relevant metric, as it represents the additional "
            "time available for ground crew deployment, evacuation, or other interventions."
        )
        add_highlighted_paragraph(doc, finding_6, new_finding_7)
    else:
        print("  WARNING: Could not find finding #6 to anchor new finding #7")

    # ─────────────────────────────────────────────────────────────────────────
    # Save
    # ─────────────────────────────────────────────────────────────────────────
    print(f"Saving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!")
    print()
    print("Summary of changes (all highlighted in YELLOW in the new document):")
    print("  1. §5.8.5: Replaced 2 paragraphs with extended-run battery analysis")
    print("  2. §5.8.5: Added 2 new paragraphs (high-wind battery + caveats)")
    print("  3. §5.8.5: Added updated battery figure from long run")
    print("  4. §5.9.2 Finding #5: Refined battery framing with concrete numbers")
    print("  5. §5.9.2 Finding #7 (NEW): Drones delay rather than contain")


if __name__ == "__main__":
    main()
