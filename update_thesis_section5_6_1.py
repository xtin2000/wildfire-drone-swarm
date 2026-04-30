#!/usr/bin/env python3
"""
update_thesis_section5_6_1.py — Insert §5.6.1 covering the earlier acoustic-only
fleet-size experiments and the rationale for fixing N=100 in subsequent work.

Builds on Thesis_04_29_Section5_Conservative.docx. Saves as new file.
"""

import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_Conservative.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_with_5_6_1.docx"


def find_paragraph_index(doc, snippet):
    for i, p in enumerate(doc.paragraphs):
        if snippet in p.text:
            return i
    return -1


def add_paragraph_after(doc, ref_para, text, style="Normal", highlight=True):
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles[style]
            if text:
                run = p.add_run(text)
                if highlight:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            return p
    return None


def main():
    print(f"Loading: {INPUT_PATH}")
    if not os.path.exists(INPUT_PATH):
        raise FileNotFoundError(INPUT_PATH)
    doc = Document(INPUT_PATH)

    # Find §5.7 Implementation Notes — §5.6.1 will go just before it
    s57 = find_paragraph_index(doc, "5.7 Implementation Notes")
    if s57 < 0:
        # try alternate heading variants
        s57 = find_paragraph_index(doc, "5.7 ")
    if s57 < 0:
        raise RuntimeError("Could not locate §5.7 anchor")

    print(f"§5.7 anchor at paragraph {s57}: '{doc.paragraphs[s57].text[:60]}'")

    # Insert before §5.7 — anchor on the paragraph just before it
    anchor = doc.paragraphs[s57 - 1]

    # Build content blocks tail-first (since each addnext puts content immediately after anchor)
    blocks = [
        ("normal",
         "Two further observations from these earlier acoustic-only experiments shaped the "
         "subsequent design. First, the 5,000-tick run length used in the fleet-size sweep "
         "was insufficient to reach natural endpoints — most scenarios were still actively "
         "burning at the cap rather than having converged to mission completion or fuel "
         "exhaustion. The subsequent experiments (§5.8 onwards) therefore extend each run to "
         "35,000 ticks (3,500 simulated seconds, approximately 58 minutes) to allow scenarios "
         "to reach natural endpoints. Second, the 100-drone fleet under acoustic suppression "
         "alone produced a delay rather than a containment — the fire was reduced in scale "
         "but continued spreading throughout the run window. This observation motivated the "
         "addition of rotor-wash physics in §5.3.5: the acoustic mechanism alone, even at the "
         "saturation fleet size, was insufficient to reach the containment threshold."),

        ("normal",
         "These results indicate a clean threshold-and-saturation behaviour in fleet-size "
         "scaling. Below 100 drones, partial coverage is not merely suboptimal — at 25 drones, "
         "outcomes were marginally worse than no drones at all, plausibly because the partial "
         "fleet's coordination overhead and the reorganisation of fire spread around the few "
         "suppressed cells produced no net improvement. Between 50 and 100 drones the "
         "configuration crosses a threshold at which the swarm's per-tick suppression capacity "
         "overtakes the fire's per-tick spread rate, producing a step change in outcome. Above "
         "100 drones the configuration saturates: additional units do not add measurable "
         "suppression because the active fire perimeter at any moment is already adequately "
         "covered. This threshold-and-saturation pattern is the principal reason the "
         "subsequent experiments in §5.6 onwards fix the fleet size at 100 drones rather than "
         "treating fleet size as an independent variable: the acoustic-only sweep had already "
         "identified 100 as the smallest fleet at which the swarm has any meaningful effect on "
         "fire dynamics, and the addition of rotor-wash physics modifies the per-drone "
         "effectiveness rather than the threshold-and-saturation structure of the fleet-size "
         "response."),

        ("normal",
         "The headline outcomes are summarised below: with 0, 25, or 50 drones the fire "
         "behaves essentially as in the no-drone control case (peak simultaneously-burning "
         "cells ≈ 1,400–1,550, total burned ≈ 4,100–4,400, no measurable improvement and in "
         "the 25-drone case a marginal worsening). At 100 drones the outcome changes "
         "qualitatively: peak simultaneously-burning cells drop to approximately 456 — a 68% "
         "reduction relative to the no-drone control — and total burned drops to approximately "
         "1,120 (a 73% reduction). At 150 drones the additional 50 drones produce no further "
         "improvement (peak 476, burned 1,121), indicating that the configuration saturates "
         "between 100 and 150 drones."),

        ("normal",
         "Five scenarios were run at 5,000 ticks (500 simulated seconds, approximately 8 "
         "minutes) under the acoustic-only physics: no_drones (0 UAVs, control), drones_25 "
         "(25 UAVs), drones_50 (50 UAVs), baseline (100 UAVs), and drones_150 (150 UAVs). All "
         "scenarios used a fixed random seed, the three-cell ignition pattern at the grid "
         "centre, and 5 m/s easterly wind."),

        ("normal",
         "Before the rotor-wash physics described in §5.3.5 was added to the simulation, an "
         "earlier set of fleet-size sweep experiments was conducted using the acoustic "
         "suppression mechanism in isolation. The aim of these earlier experiments was to "
         "identify the smallest fleet size at which a coordinated swarm produces a "
         "qualitative change in fire dynamics, so that subsequent experiments — including "
         "those reported in §5.8 — could fix the fleet size at the relevant threshold and "
         "treat the suppression mechanism rather than the fleet size as the independent "
         "variable."),

        ("heading", "5.6.1 Earlier Acoustic-Only Fleet-Size Experiments"),
    ]

    for kind, text in blocks:
        if kind == "heading":
            add_paragraph_after(doc, anchor, text, style="Heading 3")
        else:
            add_paragraph_after(doc, anchor, text)

    # Save
    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!\n")
    print("Summary:")
    print("  - §5.6.1 (NEW): Earlier Acoustic-Only Fleet-Size Experiments")
    print("    - Five scenarios at 5,000 ticks: 0/25/50/100/150 drones, acoustic-only")
    print("    - Headline result: 0/25/50 ≈ 1,400 peak / 4,100 burned;")
    print("                      100/150 ≈ 460 peak / 1,120 burned")
    print("    - Threshold-and-saturation framing: 100 drones is the smallest")
    print("      fleet that produces a qualitative change; 150 saturates it")
    print("    - Justifies fixing N=100 in §5.6 onwards")
    print("    - Notes: 5,000-tick window too short → motivated 35,000-tick re-design")
    print("    - Notes: acoustic-only delays rather than contains → motivated rotor wash addition")


if __name__ == "__main__":
    main()
