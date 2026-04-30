#!/usr/bin/env python3
"""
update_thesis_combined.py — Apply combined-mechanism findings to thesis Section 5.

Replaces sections 5.6, 5.8, and 5.9 with new content based on the rotor-wash
+ acoustic combined-mechanism experimental results. Adds new section 5.3.5
describing rotor-wash aerodynamics. All changes highlighted in yellow.

Wind-impact subsections include [PENDING] placeholders to be filled in once
the high_wind scenarios complete.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_20_ChristineLangmayr.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_27_Combined_Mechanism.docx"
FIG_DIR = "/Users/xtin2000/wildfire_drone_swarm/results_v2/figures"


# ── Helper functions ─────────────────────────────────────────────────────────

def find_paragraph_by_text(doc, snippet):
    """Return the first paragraph whose text contains the given snippet."""
    for p in doc.paragraphs:
        if snippet in p.text:
            return p
    return None


def find_paragraph_index(doc, snippet):
    """Return the index of the first paragraph whose text contains the given snippet."""
    for i, p in enumerate(doc.paragraphs):
        if snippet in p.text:
            return i
    return -1


def add_paragraph_after(doc, ref_para, text, style="Normal", highlight=True, bold=False):
    """Add a new paragraph after the reference, optionally highlighted yellow."""
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    new_para = None
    for p in doc.paragraphs:
        if p._element is new_p:
            new_para = p
            break
    if new_para is None:
        return None
    new_para.style = doc.styles[style]
    if text:
        run = new_para.add_run(text)
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if bold:
            run.bold = True
    return new_para


def replace_paragraph_text(para, new_text, highlight=True):
    """Clear all runs and replace with new text, optionally highlighted."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
        if highlight:
            para.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        run = para.add_run(new_text)
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def insert_figure_after(doc, ref_para, fig_path, caption):
    """Insert a centered figure with caption after the reference paragraph."""
    # Caption first (will end up below figure since we insert in reverse)
    cap_p = OxmlElement("w:p")
    ref_para._element.addnext(cap_p)
    for p in doc.paragraphs:
        if p._element is cap_p:
            p.style = doc.styles["Normal"]
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(10)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    # Figure
    if os.path.exists(fig_path):
        fig_p = OxmlElement("w:p")
        ref_para._element.addnext(fig_p)
        for p in doc.paragraphs:
            if p._element is fig_p:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(fig_path, width=Inches(5.5))
                break


def delete_paragraphs_between(doc, start_text_after, end_text_before):
    """Delete all paragraphs strictly between two anchor paragraphs.

    Useful for removing sections that will be replaced.
    """
    start_idx = find_paragraph_index(doc, start_text_after)
    end_idx = find_paragraph_index(doc, end_text_before)
    if start_idx < 0 or end_idx < 0 or end_idx <= start_idx:
        print(f"  WARN: couldn't find anchors '{start_text_after}' / '{end_text_before}'")
        return
    # Delete paragraphs from start_idx+1 to end_idx-1, in reverse
    for i in range(end_idx - 1, start_idx, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)


# ── Main update ──────────────────────────────────────────────────────────────

def main():
    print(f"Loading: {INPUT_PATH}")
    doc = Document(INPUT_PATH)

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE 1: §5.1 Overview — add dual-mechanism sentence
    # ─────────────────────────────────────────────────────────────────────────
    print("Updating §5.1 Overview (add dual-mechanism note)...")
    para = find_paragraph_by_text(
        doc, "swarm coordination logic that are central to this thesis"
    )
    if para:
        add_paragraph_after(
            doc, para,
            "The simulation models two physical suppression mechanisms identified in "
            "Chapter 4: directional acoustic emission from a downward-facing subwoofer (§4.2) "
            "and rotor-wash aerodynamic disturbance (§4.13). Both mechanisms are present "
            "whenever a drone hovers, reflecting the physical reality that multirotor "
            "downwash cannot be “turned off” on a flying platform.",
        )

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE 2: ADD new §5.3.5 Rotor-Wash Aerodynamic Suppression
    # Insert after §5.3.4 (Acoustic Suppression Module) content, before §5.4
    # ─────────────────────────────────────────────────────────────────────────
    print("Inserting new §5.3.5 Rotor-Wash Aerodynamic Suppression...")
    # Find the last paragraph of §5.3.4 (just before §5.4)
    sec54_idx = find_paragraph_index(doc, "5.4 Swarm Coordination")
    if sec54_idx > 0:
        # Insert before §5.4 heading: place anchored on the paragraph just before it
        anchor = doc.paragraphs[sec54_idx - 1]

        # Build content blocks IN REVERSE ORDER (since addnext inserts immediately after)
        blocks = [
            ("normal", "Rotor-wash effects are active whenever a drone is in the HOVERING or "
                       "EMITTING state. Drones in TRANSIT, RETURNING, or CHARGING produce no "
                       "rotor wash effect on suppression because they are not holding station "
                       "above a target."),
            ("normal", "Embers lofted by burning cells are tracked as ballistic particles (§5.2.4). "
                       "When an ember passes through the rotor wash column of any hovering drone, "
                       "an additional downward acceleration is applied (1.5 m/s² peak, scaled "
                       "by radial position and altitude decay). This shortens the ember’s "
                       "airborne trajectory and reduces long-range spotting downwind."),
            ("normal", "When a drone hovers above a STATE_EMBER cell — a small, nascent "
                       "ignition with low burn intensity — the rotor wash deposits suppression "
                       "energy at a rate scaled to local downwash strength. Embers are easier to "
                       "extinguish than fully burning cells (consistent with §4.13’s "
                       "observation that lightweight embers are readily deflected by moderate "
                       "airflow), so the ember suppression rate (50 J/s at peak wash strength) is "
                       "an order of magnitude higher than the rate applied to established burning "
                       "cells (5 J/s). This asymmetry reflects the physical reality that rotor "
                       "wash collapses small flamelets effectively but cannot overcome the "
                       "convective updraft of an established fire."),
            ("normal", "Using actuator-disk theory, the induced velocity at the rotor disk for a "
                       "drone of mass m hovering in air of density ρ with total rotor area A is "
                       "v_disk = √(m·g / (2·ρ·A)). For the modeled 5 kg "
                       "quadrotor with 0.20 m² total disk area in standard air density, this "
                       "produces v_disk ≈ 11.2 m/s at the rotor. Below the disk, the wake "
                       "decays linearly with altitude due to ambient air entrainment, modeled with "
                       "a 10 m decay length. The ground-level downwash footprint expands with "
                       "altitude, beginning at 1.5 m radius directly under the drone and spreading "
                       "at 0.15 m radius per meter of altitude."),
            ("normal", "The second suppression mechanism is the downward airflow column produced "
                       "by the drone’s rotors. Unlike the acoustic emitter, rotor wash is "
                       "always present whenever the drone is hovering or holding station — "
                       "it is a physical consequence of generating lift and cannot be disabled. "
                       "The model captures three effects: a downwash velocity field below the "
                       "drone, direct disruption of nascent ember-state ignitions, and "
                       "aerodynamic deflection of airborne embers in flight."),
            ("heading", "5.3.5 Rotor-Wash Aerodynamic Suppression"),
        ]

        for kind, text in blocks:
            if kind == "heading":
                add_paragraph_after(doc, anchor, text, style="Heading 3")
            else:
                add_paragraph_after(doc, anchor, text)

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE 3: §5.6 Experimental Scenarios — REPLACE
    # ─────────────────────────────────────────────────────────────────────────
    print("Replacing §5.6 Experimental Scenarios...")
    # The previous update added Table 1 and 7 scenarios. Now we replace with 5 scenarios.
    # Find anchor: §5.6 heading
    sec56_idx = find_paragraph_index(doc, "5.6 Experimental Scenarios")
    sec57_idx = find_paragraph_index(doc, "5.7 Implementation Notes")

    if sec56_idx >= 0 and sec57_idx > sec56_idx:
        # Delete everything between heading and §5.7
        for i in range(sec57_idx - 1, sec56_idx, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

        # Insert new content after the §5.6 heading
        anchor = doc.paragraphs[sec56_idx]
        blocks_56 = [
            ("normal", "All scenarios share the same three-cell ignition cluster at the grid "
                       "center and use a fixed random seed (42) for reproducibility. Each scenario "
                       "terminates when no burning or ember-state cells remain (mission success) or "
                       "when 35,000 ticks elapse, whichever comes first."),
            ("normal", "The Rotor-Wash-Only scenario is included not because acoustic emission "
                       "could realistically be disabled on a flying drone — it cannot — "
                       "but because the comparison between Rotor-Wash-Only and Combined isolates "
                       "the additional contribution of acoustic suppression beyond the rotor wash "
                       "that is always present. This is the central experimental question: what "
                       "does the acoustic emitter add to a swarm whose drones already produce "
                       "downwash by virtue of flying?"),
            ("normal", "Five scenarios are defined: (1) No Drones — 0 drones, 5 m/s wind, "
                       "counterfactual baseline; (2) Rotor Wash Only — 100 drones, 5 m/s wind, "
                       "rotor wash active but acoustic emission disabled; (3) Combined — 100 "
                       "drones, 5 m/s wind, both mechanisms active (realistic full-physics drone); "
                       "(4) Rotor Wash + High Wind — 100 drones, 10 m/s SE wind, rotor only; "
                       "(5) Combined + High Wind — 100 drones, 10 m/s SE wind, both "
                       "mechanisms."),
            ("normal", "The experimental design for evaluating the simulation framework was "
                       "restructured around five scenarios that isolate the contribution of each "
                       "suppression mechanism. The previous design used acoustic suppression in "
                       "isolation and ran for 5,000 ticks (500 simulated seconds); the revised "
                       "design adds rotor-wash physics (always present on a hovering drone) and "
                       "extends each scenario to 35,000 ticks (3,500 simulated seconds, "
                       "approximately 58 minutes of mission time) to allow scenarios to reach "
                       "natural endpoints — either mission completion or fuel exhaustion."),
        ]
        for _, text in blocks_56:
            add_paragraph_after(doc, anchor, text)

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE 4: §5.8 Simulation Results — REPLACE ENTIRELY
    # ─────────────────────────────────────────────────────────────────────────
    print("Replacing §5.8 Simulation Results...")
    sec58_idx = find_paragraph_index(doc, "5.8 Simulation Results")
    sec59_idx = find_paragraph_index(doc, "5.9 Analysis and Discussion")

    if sec58_idx >= 0 and sec59_idx > sec58_idx:
        # Delete everything between
        for i in range(sec59_idx - 1, sec58_idx, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

        anchor = doc.paragraphs[sec58_idx]

        # Build content (reverse order for insertion)
        blocks_58 = []

        # Battery section (5.8.6)
        blocks_58.extend([
            ("normal", "The implication for operational deployment is encouraging: when the swarm "
                       "is effective enough to contain the fire (combined mechanism), the "
                       "engagement is brief and battery endurance is comfortable. When the swarm "
                       "cannot contain the fire, battery becomes the limiting factor for mission "
                       "duration — but in those cases, the operational value of continuing "
                       "the engagement is questionable anyway."),
            ("normal", "Combined baseline: mission complete at t ≈ 1,000 s with mean battery "
                       "around 93%. Battery is not a limiting factor for short successful "
                       "missions. Rotor-wash-only baseline: ran the full 35,000 ticks with mean "
                       "battery declining to approximately 60%. Sustained operation against an "
                       "uncontainable fire continues to deplete batteries linearly until "
                       "intervention ends."),
            ("normal", "Battery dynamics differ markedly between the scenarios that achieve "
                       "containment and those that do not."),
            ("heading", "5.8.6 Battery and Operational Endurance"),
        ])

        # Wind impact section (5.8.5) — PENDING
        blocks_58.extend([
            ("normal", "[PENDING: Once rotor_only_high_wind and combined_high_wind scenarios "
                       "complete, this subsection will compare their outcomes to the baseline "
                       "scenarios and discuss whether the combined mechanism's containment "
                       "capability survives 10 m/s wind. Early data shows drone stability "
                       "collapses to near-zero (≈0.02) under high wind, suggesting "
                       "containment may not be possible regardless of mechanism.]"),
            ("heading", "5.8.5 Wind Impact"),
        ])

        # Mechanism decomposition (5.8.4)
        blocks_58.extend([
            ("normal", "The combined mechanism's effectiveness is greater than the sum of its "
                       "parts. Rotor wash alone reduces peak burning by 19% (3,930 → 3,185) "
                       "and peak embers by 44% (393 → 219). The combined effect — 97% "
                       "reduction in peak burning and 99% reduction in embers — is "
                       "qualitatively different. This is consistent with the operational reading: "
                       "acoustic and aerodynamic mechanisms target different fire stages (flame "
                       "front vs ember layer), and a swarm that can suppress only one of them is "
                       "overwhelmed by the other."),
            ("normal", "Comparing across the three baseline-wind scenarios isolates the "
                       "contribution of each mechanism to fire containment. Rotor wash alone "
                       "reduces peak burning modestly (19%) but reduces peak embers substantially "
                       "(44%), consistent with §4.13's argument that aerodynamic disruption is "
                       "most effective on ember-scale ignitions. Acoustic emission (when added "
                       "to rotor wash) extends suppression to established flame fronts, enabling "
                       "the swarm to suppress fire faster than it spreads."),
            ("heading", "5.8.4 Mechanism Decomposition"),
        ])

        # Combined Performance figure
        blocks_58.append(
            ("figure", f"{FIG_DIR}/fig_total_fire.png",
             "Figure: Total active fire cells (burning + ember) across scenarios. The combined "
             "mechanism (blue) maintains near-zero active fire throughout, while no_drones "
             "(orange), rotor_only_baseline (green), and rotor_only_high_wind (red) all "
             "reach high active-fire counts before the fire eventually burns out by fuel "
             "exhaustion.")
        )

        # Combined Performance section (5.8.3) — the headline
        blocks_58.extend([
            ("normal", "The mission-complete result at 16.6 minutes is a strong claim, but it is "
                       "bounded by the simulation’s idealized assumptions (§5.9.3): perfect "
                       "coordination, idealized acoustic propagation, no thermal damage to "
                       "drones, flat homogeneous terrain. The headline finding should therefore "
                       "be read not as a literal performance estimate but as evidence that the "
                       "combined mechanism crosses a containment threshold that neither "
                       "mechanism alone can reach."),
            ("normal", "The two mechanisms are complementary in exactly the way Chapter 4 "
                       "predicted. Acoustic emission disrupts established burning cells, "
                       "reducing intensity and extinguishing them with sufficient sustained "
                       "exposure. Rotor wash handles the ember layer — both ground-level "
                       "ember-state cells and airborne embers in flight. Neither mechanism "
                       "alone can do both jobs adequately; together they cover the full fire "
                       "lifecycle from ignition to flame-front to spotting."),
            ("normal", "When both mechanisms are active, the swarm transforms from a delaying "
                       "agent into a containing one. Peak burning was only 108 cells (a 97% "
                       "reduction vs no drones), total burned was 1,388 cells (a 96% reduction), "
                       "peak embers was just 3 cells (a 99% reduction), and mission completed at "
                       "t = 999 s (16.6 simulated minutes)."),
            ("heading", "5.8.3 Combined Acoustic + Rotor-Wash Performance"),
        ])

        # Rotor-wash-only section (5.8.2)
        blocks_58.extend([
            ("normal", "This isolated rotor-wash performance is not operationally proposed "
                       "— drones cannot fly without producing rotor wash — but it "
                       "usefully bounds the contribution of aerodynamic suppression alone. The "
                       "44% ember reduction is the standout result: drones hovering above ember "
                       "cells extinguish them efficiently, and downwash on airborne embers "
                       "shortens their trajectories."),
            ("normal", "The pattern is consistent with the physical predictions of §4.13: rotor "
                       "wash is effective at ember management but cannot meaningfully suppress "
                       "an established flame front. With acoustic emission disabled, the swarm "
                       "has no mechanism to reduce burn intensity on cells that have already "
                       "established sustained combustion, and the fire continues to consume the "
                       "available fuel until natural burnout."),
            ("normal", "With rotor wash active but acoustic emission disabled, the 100-drone "
                       "swarm produced measurable but modest improvements over the no-drone "
                       "baseline. Peak burning was reduced from 3,930 to 3,185 cells (a 19% "
                       "reduction), total burned was reduced from 37,606 to 37,525 cells "
                       "(essentially unchanged), and peak embers were reduced from 393 to 219 "
                       "cells (a 44% reduction)."),
            ("heading", "5.8.2 Rotor-Wash-Only Performance"),
        ])

        # No-drones control (5.8.1)
        blocks_58.extend([
            ("normal", "The 10 ground firefighters operating throughout the scenario applied "
                       "water to 1,976 cells (4.9% of grid) but were unable to materially affect "
                       "the spread rate at the modeled walking speed of 0.5 m/s."),
            ("normal", "The no-drone scenario establishes the counterfactual: in the absence of "
                       "any UAV intervention, the fire spread monotonically from the ignition "
                       "cluster, reaching peak burning of 3,930 cells at approximately t = 900 s "
                       "and ultimately consuming 37,606 cells. The fire extinguished naturally "
                       "at approximately t = 1,500 s, having exhausted the available unburned "
                       "fuel adjacent to the burned area. Peak ember activity reached 393 cells, "
                       "reflecting the full ember-lofting potential of an unsuppressed fire."),
            ("heading", "5.8.1 No-Drone Control"),
        ])

        # Section intro (with summary table)
        blocks_58.extend([
            ("normal", "The headline finding is dramatic: the combined-mechanism swarm fully "
                       "contains the fire in 16.6 simulated minutes, with only 3.5% of the grid "
                       "(1,388 of 40,000 cells) burned. Neither the no-drone control nor the "
                       "rotor-wash-only configuration achieved containment within 35,000 ticks; "
                       "both allowed the fire to consume approximately 94% of the available grid."),
            ("normal", "Headline outcomes at t = 3,500 s (or mission completion): "
                       "(a) No Drones — peak burning 3,930, total burned 37,606, peak embers "
                       "393, no containment. "
                       "(b) Rotor Wash Only — peak burning 3,185, total burned 37,525, peak "
                       "embers 219, no containment. "
                       "(c) Combined — peak burning 108, total burned 1,388, peak embers 3, "
                       "MISSION COMPLETE at t = 999 s (16.6 min). "
                       "(d) Rotor + High Wind — [PENDING]. "
                       "(e) Combined + High Wind — [PENDING]."),
            ("normal", "The five scenarios were executed headlessly. The headline results are "
                       "presented below; subsequent subsections analyze each scenario in detail."),
        ])

        # Insert all in order (we list them tail-first so reverse iteration places them top-first)
        for block in blocks_58:
            kind = block[0]
            if kind == "heading":
                add_paragraph_after(doc, anchor, block[1], style="Heading 3")
            elif kind == "figure":
                insert_figure_after(doc, anchor, block[1], block[2])
            else:
                add_paragraph_after(doc, anchor, block[1])

    # ─────────────────────────────────────────────────────────────────────────
    # CHANGE 5: §5.9.2 Key Findings — REPLACE
    # ─────────────────────────────────────────────────────────────────────────
    print("Replacing §5.9.2 Key Findings...")
    f592_idx = find_paragraph_index(doc, "5.9.2 Key Findings")
    f593_idx = find_paragraph_index(doc, "5.9.3 Limitations")

    if f592_idx >= 0 and f593_idx > f592_idx:
        for i in range(f593_idx - 1, f592_idx, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

        anchor = doc.paragraphs[f592_idx]

        new_findings = [
            "6. The simulation results are bounded by its idealized assumptions, but the "
            "direction of the effect is robust. The mission-complete result should not be read as "
            "a literal claim that 100 drones could extinguish a real grass fire in 17 minutes; it "
            "is, rather, evidence that crossing the suppression-pressure threshold from "
            "“delay” to “contain” is possible only when both mechanisms are "
            "active. The qualitative finding — that acoustic + rotor-wash combined produces "
            "fundamentally different dynamics than either alone — is robust to reasonable "
            "variation in parameter calibration because it depends on the structural "
            "complementarity of the two mechanisms, not on the absolute magnitudes.",

            "5. Battery endurance is not the binding constraint when the swarm is effective. The "
            "combined-mechanism mission completed in 17 minutes with mean fleet battery at 93%, "
            "well above the 20% return-to-base threshold. Battery becomes a constraint only in "
            "scenarios where the swarm cannot contain the fire and continues hovering "
            "indefinitely — situations in which the operational value of continued "
            "engagement is itself questionable.",

            "4. Wind impact on the combined mechanism. [PENDING — to be filled in once "
            "high_wind scenarios complete. Early data shows drone stability collapses to near-zero "
            "under 10 m/s wind, suggesting containment may not survive high-wind conditions "
            "regardless of mechanism.]",

            "3. Ember management is the highest-value contribution per drone. Peak ember counts "
            "dropped from 393 (no drones) to 219 (rotor wash only) to 3 (combined). The 99% "
            "reduction in the combined scenario directly reflects the swarm’s ability to "
            "extinguish nascent ignitions before they establish sustained combustion. This "
            "empirically supports §4.14’s redirection of UAV suppression toward "
            "early-ignition and ember-management roles rather than frontline flame suppression.",

            "2. The two suppression mechanisms are complementary rather than redundant. Rotor "
            "wash alone reduces peak burning by only 19% but reduces peak embers by 44%, "
            "consistent with §4.13’s argument that aerodynamic disruption is most effective "
            "on ember-scale ignitions. Acoustic suppression is most effective on burning cells. "
            "Together, the two mechanisms cover the full fire lifecycle and produce qualitatively "
            "different containment behavior than either alone.",

            "1. The combination of acoustic and rotor-wash mechanisms can fully contain a "
            "coordinated wildfire under moderate wind conditions. A 100-drone swarm with both "
            "mechanisms active extinguished the modeled fire in 16.6 simulated minutes, with only "
            "3.5% of the grid burned. Neither mechanism alone achieved containment.",

            "The simulation results support the following conclusions:",
        ]

        for text in new_findings:
            add_paragraph_after(doc, anchor, text)

    # ─────────────────────────────────────────────────────────────────────────
    # Save
    # ─────────────────────────────────────────────────────────────────────────
    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!\n")
    print("Summary of changes (all yellow-highlighted in the new document):")
    print("  - §5.1: Added dual-mechanism note")
    print("  - §5.3.5 (NEW): Rotor-Wash Aerodynamic Suppression subsection")
    print("  - §5.6: Replaced with new 5-scenario design (mechanism comparison)")
    print("  - §5.8: Entirely rewritten with new combined-mechanism findings")
    print("    - 5.8.1 No-Drone Control")
    print("    - 5.8.2 Rotor-Wash-Only Performance")
    print("    - 5.8.3 Combined Performance (THE HEADLINE)")
    print("    - 5.8.4 Mechanism Decomposition")
    print("    - 5.8.5 Wind Impact [PENDING]")
    print("    - 5.8.6 Battery and Operational Endurance")
    print("  - §5.9.2: Rewritten Key Findings (6 findings, finding #4 PENDING)")


if __name__ == "__main__":
    main()
