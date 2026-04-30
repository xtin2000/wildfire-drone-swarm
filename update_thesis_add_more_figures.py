#!/usr/bin/env python3
"""Insert three additional figures into the thesis:
  - fig_active_fire.png       at end of §5.8.3 Combined Performance
  - fig_suppression_pressure.png at end of §5.8.4 Decomposition
  - fig_battery_depletion.png  at end of §5.8.6 Battery and Operational Endurance

Builds on Thesis_04_29_Section5_v6.docx. Saves as new file.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_v6.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_v7.docx"
FIG_DIR = "/Users/xtin2000/wildfire_drone_swarm/results_v2/figures"

CAPTION_ACTIVE_FIRE = (
    "Figure: Total active fire (burning + ember cells) over simulated time across the major "
    "scenarios at 5 m/s wind (combined_high_wind shown dashed at 10 m/s for contrast). The "
    "combined_baseline trajectory (green) collapses to zero within ~17 simulated minutes; "
    "drones_only (orange) and rotor_only_baseline (blue) climb into the thousands and run "
    "the full window; the no-intervention scenarios (pure_fire, no_drones) follow the "
    "uncontained trajectory. The active-fire view (rather than burning-only) emphasizes "
    "the ember-management contribution of the combined mechanism: peak active fire is "
    "two orders of magnitude lower under the combined configuration than under any other."
)

CAPTION_SUPPRESSION_PRESSURE = (
    "Figure: Suppression pressure — ratio of active drones (emitting + hovering) to active "
    "fire cells (burning + ember) over simulated time, log scale. A ratio above 1.0 (parity, "
    "shown dotted) means the swarm has at least one drone holding station against every "
    "active fire cell. The combined_baseline configuration (green) climbs above parity "
    "within the first minute and stays there; rotor_only_baseline (blue) falls below parity "
    "within the first 5 minutes and never recovers; drones_only (orange) starts above "
    "parity but loses ground as the fire spreads in two directions and the swarm cannot "
    "cover both flanks. The suppression-pressure metric directly visualises the early-spread "
    "race that determines containment outcome (§5.8.7)."
)

CAPTION_BATTERY = (
    "Figure: Drone fleet battery depletion over simulated time. Solid lines show mean "
    "battery across the fleet; dashed lines show the minimum battery (lowest charge across "
    "all drones). For the contained combined_baseline (green), mean battery drains less "
    "than 7 percentage points before mission_complete fires at ~17 simulated minutes — "
    "battery is comfortably above the 20% RTB threshold throughout. For the uncontained "
    "rotor_only_baseline (blue), mean battery declines linearly across the full 35,000-tick "
    "window to approximately 60%. For the high-wind scenario (red), frequent excursions "
    "above the 12 m/s safety threshold produce stepwise battery depletion as drones cycle "
    "through return-to-base and recharge phases. Battery is therefore the binding constraint "
    "only in scenarios where the swarm cannot contain the fire."
)


def find_paragraph_index(doc, snippet):
    for i, p in enumerate(doc.paragraphs):
        if snippet in p.text:
            return i
    return -1


def insert_figure_after(doc, ref_para, fig_path: str, caption: str) -> None:
    """Insert centered figure followed by centered yellow-highlighted italic caption."""
    cap_p = OxmlElement("w:p")
    ref_para._element.addnext(cap_p)
    for p in doc.paragraphs:
        if p._element is cap_p:
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(10)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            break

    if os.path.exists(fig_path):
        fig_p = OxmlElement("w:p")
        ref_para._element.addnext(fig_p)
        for p in doc.paragraphs:
            if p._element is fig_p:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(fig_path, width=Inches(5.5))
                break
    else:
        print(f"WARN: figure not found: {fig_path}")


def insert_at_end_of_section(doc, current_anchor, next_anchor, fig_path, caption):
    """Insert a figure at the end of a subsection (just before the next subsection's heading)."""
    s_idx = find_paragraph_index(doc, current_anchor)
    e_idx = find_paragraph_index(doc, next_anchor)
    if s_idx < 0 or e_idx < 0:
        raise RuntimeError(f"Anchors not found: '{current_anchor}' / '{next_anchor}'")
    print(f"  current section heading at {s_idx}, next at {e_idx}")
    # Insert before next-section heading: anchor on paragraph immediately before it
    anchor_para = doc.paragraphs[e_idx - 1]
    insert_figure_after(doc, anchor_para, fig_path, caption)


def main() -> None:
    print(f"Loading: {INPUT_PATH}")
    if not os.path.exists(INPUT_PATH):
        raise FileNotFoundError(INPUT_PATH)
    doc = Document(INPUT_PATH)

    # ── §5.8.3 Combined Performance — insert active-fire figure at end ──
    print("Inserting active-fire figure at end of §5.8.3...")
    insert_at_end_of_section(
        doc,
        "5.8.3 Combined Acoustic + Rotor-Wash + Ground-Crew Performance",
        "5.8.4 Decomposition: Mechanisms and Agents",
        f"{FIG_DIR}/fig_active_fire.png",
        CAPTION_ACTIVE_FIRE,
    )

    # ── §5.8.4 Decomposition — insert suppression-pressure figure at end ──
    print("Inserting suppression-pressure figure at end of §5.8.4...")
    insert_at_end_of_section(
        doc,
        "5.8.4 Decomposition: Mechanisms and Agents",
        "5.8.5 Wind Impact",
        f"{FIG_DIR}/fig_suppression_pressure.png",
        CAPTION_SUPPRESSION_PRESSURE,
    )

    # ── §5.8.6 Battery — insert battery-depletion figure at end ──
    print("Inserting battery-depletion figure at end of §5.8.6...")
    insert_at_end_of_section(
        doc,
        "5.8.6 Battery and Operational Endurance",
        "5.8.7 Sensitivity Analysis",
        f"{FIG_DIR}/fig_battery_depletion.png",
        CAPTION_BATTERY,
    )

    # Save
    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!\n")
    print("Summary:")
    print("  - §5.8.3: + fig_active_fire.png (total active fire over time, all scenarios)")
    print("  - §5.8.4: + fig_suppression_pressure.png (drones per active fire cell, log scale)")
    print("  - §5.8.6: + fig_battery_depletion.png (mean+min battery, 3 scenarios)")


if __name__ == "__main__":
    main()
