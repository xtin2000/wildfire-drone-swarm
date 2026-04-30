#!/usr/bin/env python3
"""Replace the 2x2 factorial figure in the thesis with the version showing
the min/max range bar on the combined cell, and update the caption text.

Builds on Thesis_04_29_Section5_with_5_6_1.docx. Saves as new file.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_with_5_6_1.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_v5.docx"
FIG_PATH = "/Users/xtin2000/wildfire_drone_swarm/results_v2/figures/fig_2x2_factorial.png"

NEW_CAPTION = (
    "Figure: 2×2 factorial — % of grid burned at the four corners of the drones × firefighters "
    "design (5 m/s wind, Option B firefighter physics). The combined cell aggregates over N=5 "
    "random-seed trials of combined_baseline (seeds 42-46): the bar shows the median outcome "
    "(3.5%); the vertical range bar spans the observed minimum (0.1%) and maximum (93.4%) and "
    "annotates the bimodal nature of the outcome — 3 of 5 trials achieved containment, 2 of 5 "
    "lost the early-spread race and ran to fuel exhaustion. The other three corners are "
    "essentially seed-insensitive at this scale (all run to ~93–99% regardless). The qualitative "
    "finding — only the joint configuration crosses the containment threshold — is unchanged."
)


def main() -> None:
    print(f"Loading: {INPUT_PATH}")
    if not os.path.exists(INPUT_PATH):
        raise FileNotFoundError(INPUT_PATH)
    doc = Document(INPUT_PATH)

    # Find the existing caption paragraph (it contains "2×2 factorial — % of grid burned")
    cap_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "2×2 factorial" in p.text and "% of grid burned" in p.text:
            cap_idx = i
            break
    if cap_idx < 0:
        raise RuntimeError("Could not find existing 2×2 caption paragraph")

    print(f"Found caption at paragraph {cap_idx}")

    # The figure paragraph is immediately before the caption paragraph
    fig_idx = cap_idx - 1
    fig_para = doc.paragraphs[fig_idx]
    cap_para = doc.paragraphs[cap_idx]

    # The previous insertion put the figure immediately before the caption, so we trust
    # that fig_idx = cap_idx - 1 is the figure paragraph.

    # Anchor for inserts: paragraph just before the figure
    anchor_idx = fig_idx - 1
    anchor = doc.paragraphs[anchor_idx]

    # Remove old caption first (so we can rebuild cleanly), then old figure
    cap_para._element.getparent().remove(cap_para._element)
    fig_para._element.getparent().remove(fig_para._element)

    # Insert new caption first, then new figure (we use addnext, so reverse order)
    new_cap_p = OxmlElement("w:p")
    anchor._element.addnext(new_cap_p)
    for p in doc.paragraphs:
        if p._element is new_cap_p:
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(NEW_CAPTION)
            run.italic = True
            run.font.size = Pt(10)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            break

    new_fig_p = OxmlElement("w:p")
    anchor._element.addnext(new_fig_p)
    for p in doc.paragraphs:
        if p._element is new_fig_p:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(FIG_PATH, width=Inches(5.5))
            break

    # Save
    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!\n")
    print("Summary:")
    print("  - Replaced 2×2 factorial figure with version showing min/max range bar on")
    print("    the combined cell (N=5 seeds: 3 contained, 2 fuel-exhausted)")
    print("  - Updated caption to acknowledge probabilistic outcome and bimodal range")


if __name__ == "__main__":
    main()
