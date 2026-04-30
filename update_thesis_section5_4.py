#!/usr/bin/env python3
"""Rewrite §5.4 Swarm Coordination with code-grounded algorithm details.

Replaces the prose under each of §5.4, §5.4.1, §5.4.2, §5.4.3 while preserving
the heading paragraphs themselves. Builds on Thesis_04_29_Section5_v8.docx.
"""

import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_v8.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_30_Section5_v9.docx"


# ── New content for each subsection ──────────────────────────────────────────

INTRO_5_4 = [
    "Coordination of the 100-drone fleet is performed by three cooperating components, "
    "dispatched together from SwarmCoordinator.step every COORD_INTERVAL = 20 simulation "
    "ticks (2.0 simulated seconds). The mission planner first surveys grid state and emits a "
    "ranked target list (§5.4.1); the task allocator solves an optimal-assignment problem "
    "mapping drones to targets (§5.4.2); the conflict avoider produces velocity adjustments "
    "to prevent inter-drone collisions and to keep drones clear of ground firefighters "
    "(§5.4.3). The 2.0-second cycle balances responsiveness to fast-changing fire "
    "geometry against the cost of full-fleet replanning at each tick.",
]

CONTENT_5_4_1 = [
    "The mission planner enumerates every cell in STATE_BURNING or STATE_EMBER and computes "
    "a composite priority score for each, retaining the top max_targets = 200 cells (sorted "
    "descending) as the candidate task pool. The score function is:",

    "score(cell) = 10·is_ember + 5·(1 − intensity)·is_burning + 2·n_burning_neighbours + "
    "3/min_ff_dist + 0.5·|wind|",

    "where is_ember and is_burning are 0/1 state indicators, intensity is the cell's burn "
    "intensity in [0, 1], n_burning_neighbours counts how many of the eight surrounding "
    "cells are themselves in STATE_BURNING, min_ff_dist is the world-space distance from "
    "the cell to the nearest ground firefighter, and |wind| is the local wind-vector "
    "magnitude in m/s.",

    "Each weight encodes an operational priority. The 10× weight on ember-state cells "
    "reflects that small nascent ignitions are both the easiest cells to extinguish (low "
    "intensity, low convective updraft to resist external suppression) and the most "
    "valuable to kill early — once an ember-state cell grows past intensity 0.35 it "
    "transitions to STATE_BURNING and becomes an order of magnitude harder to extinguish. "
    "The 5·(1 − intensity) term applies to burning cells and prioritises low-intensity "
    "(recently ignited) cells over fully-developed flame fronts, reflecting the same "
    "ease-of-extinction asymmetry. The 2× expansion-pressure term grows with the number of "
    "burning 8-neighbours: a cell surrounded by burning cells is about to spread further "
    "and is therefore high-priority. The 3/min_ff_dist term escalates priority for cells "
    "near ground crews, capturing personnel-safety risk. The 0.5·|wind| term gives a modest "
    "boost to cells in high-wind regions where fire spread will be fastest.",

    "The top-200 list is wrapped as Target dataclass instances, each with a unique task_id, "
    "(col, row) cell coordinates, urgency equal to the score, and world coordinates derived "
    "from the cell index, and is passed to the task allocator.",
]

CONTENT_5_4_2 = [
    "Despite the historical class name AuctionTaskAllocator (which originated from an "
    "earlier bidding-style implementation), the production code solves an optimal-assignment "
    "problem rather than running a greedy bid auction. For each (drone, target) pair "
    "considered in the current cycle, the allocator computes a per-pair score:",

    "score(drone, target) = 20/distance + 5·urgency + 3·battery_fraction + 2·stability",

    "where distance is the planar (x, y) distance from the drone to the target's world "
    "position (floored at 1 m to prevent the 20/distance term from blowing up), urgency is "
    "the target's prioritisation score from §5.4.1, battery_fraction is the drone's "
    "remaining charge in [0, 1], and stability is the drone's wind-stability factor "
    "(quadratic falloff with wind speed, also in [0, 1]). The four terms balance proximity, "
    "target importance, energy reserve, and platform fitness.",

    "The cost matrix is constructed as the negation of these scores and solved using "
    "scipy.optimize.linear_sum_assignment, which implements the Hungarian algorithm "
    "(Kuhn–Munkres). The Hungarian algorithm finds the one-to-one drone-to-target mapping "
    "that minimises total cost (equivalently, maximises total swarm utility) in O(n³) time. "
    "Unlike a greedy assignment — which would assign the highest-scoring (drone, target) "
    "pair, then the next, and so on, potentially leaving high-utility drones with "
    "low-utility targets — the Hungarian assignment is provably optimal at the cycle level.",

    "Two stability mechanisms prevent oscillatory behaviour. First, drones currently in "
    "states RETURNING or CHARGING are excluded from the candidate pool — they cannot accept "
    "new tasks until they have recharged past BATTERY_FULL_THRESH = 0.90. Second, a "
    "hysteresis parameter TASK_REALLOC_HYSTERESIS = 0.1 prevents reassignment churn: a drone "
    "is moved off its existing task only if the newly-computed score exceeds its previous "
    "score by at least 0.1, otherwise the previous assignment is retained even if the "
    "Hungarian solver proposes a swap. A separate constraint, applied at the drone-agent "
    "level rather than the allocator, prevents drones in STATE_EMITTING from being "
    "interrupted to begin a new task — this avoids halting active suppression on a target "
    "that is about to be extinguished.",
]

CONTENT_5_4_3 = [
    "Inter-drone collision avoidance uses a simplified Velocity Obstacle (VO) algorithm "
    "with ORCA-inspired symmetric responsibility. The conflict avoider is given the full "
    "fleet's positions and velocities each cycle and queries scipy.spatial.KDTree for all "
    "drone pairs within 4 × MIN_DRONE_SEPARATION = 20 m of each other. For each candidate "
    "pair (i, j) the algorithm computes:",

    "Time to closest approach:  t_ca = clip( −(p_rel · v_rel) / |v_rel|², 0, "
    "VO_TIME_HORIZON = 3 s )",

    "Predicted closest distance:  d_min = | p_rel + v_rel · t_ca |",

    "where p_rel = p_j − p_i and v_rel = v_j − v_i are the relative position and velocity "
    "vectors in the horizontal plane. If the predicted closest distance d_min is less than "
    "MIN_DRONE_SEPARATION = 5 m, both drones receive a repulsion impulse directed along the "
    "relative-position vector, with magnitude proportional to (5 − d_min)/5. The two drones "
    "share responsibility 50/50 (i moves away from j by half the impulse, j moves away from "
    "i by the other half), in the spirit of ORCA's reciprocal-avoidance assumption. The "
    "impulses are stored as velocity hints and applied as additive offsets in each drone's "
    "flight controller on the next tick, in superposition with the PID position-tracking "
    "output.",

    "A separate firefighter-exclusion pass operates on the same fleet positions: each "
    "drone within FIREFIGHTER_EXCL_RADIUS = 15 m of any firefighter receives a repulsion "
    "impulse of strength up to 3.0 directed away from the firefighter. This prevents the "
    "swarm from operating directly above ground crews — a hard real-world safety "
    "constraint that protects personnel from drone failure modes (motor loss, control "
    "fault, low-altitude collision recovery).",

    "Collision avoidance handles drone-drone and drone-firefighter conflicts only. The "
    "minimum drone-to-active-fire standoff is enforced separately, in the drone agent's "
    "_safe_hover_pos geometry: when a drone is assigned a target cell, its hover position "
    "is computed at horizontal distance (SAFE_FIRE_DISTANCE + MAX_FIRE_DISTANCE)/2 from "
    "the target, set to 7.5 m under the current configuration. Thermal-damage prevention "
    "is therefore implicit in the assignment geometry rather than enforced by an explicit "
    "exclusion zone.",
]


def find_paragraph_index(doc, snippet, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if snippet in doc.paragraphs[i].text:
            return i
    return -1


def add_paragraph_after(doc, ref_para, text, style="Normal", highlight=True):
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = doc.styles[style]
            if text:
                run = p.add_run(text)
                if highlight:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            return p
    return None


def replace_section_content(doc, start_heading_snippet: str, next_heading_snippet: str,
                            new_paragraphs_reverse_order: list[str]) -> None:
    """Delete paragraphs strictly between the start and next headings, then
    insert new paragraphs immediately after the start heading (reverse order
    because addnext reverses)."""
    s_idx = find_paragraph_index(doc, start_heading_snippet)
    e_idx = find_paragraph_index(doc, next_heading_snippet, start=s_idx + 1)
    if s_idx < 0 or e_idx < 0:
        raise RuntimeError(f"Anchors not found: '{start_heading_snippet}' / '{next_heading_snippet}'")
    print(f"  start at {s_idx}, next at {e_idx}; deleting {e_idx - s_idx - 1} paragraphs in between")

    for i in range(e_idx - 1, s_idx, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)

    anchor = doc.paragraphs[s_idx]
    for text in new_paragraphs_reverse_order:
        add_paragraph_after(doc, anchor, text)


def main() -> None:
    print(f"Loading: {INPUT_PATH}")
    if not os.path.exists(INPUT_PATH):
        raise FileNotFoundError(INPUT_PATH)
    doc = Document(INPUT_PATH)

    print("Replacing §5.4 intro (between '5.4 Swarm Coordination' and '5.4.1')...")
    replace_section_content(
        doc, "5.4 Swarm Coordination", "5.4.1 Target Prioritization",
        list(reversed(INTRO_5_4)),
    )

    print("Replacing §5.4.1 content...")
    replace_section_content(
        doc, "5.4.1 Target Prioritization", "5.4.2 Task Allocation",
        list(reversed(CONTENT_5_4_1)),
    )

    print("Replacing §5.4.2 content...")
    replace_section_content(
        doc, "5.4.2 Task Allocation", "5.4.3 Collision Avoidance",
        list(reversed(CONTENT_5_4_2)),
    )

    # Find what comes after §5.4.3. Search for §5.5 heading variants.
    next_after_543_candidates = [
        "5.5 Ground Firefighter Model",
        "5.5 Firefighter Model",
        "5.5 ",
    ]
    next_anchor = None
    for cand in next_after_543_candidates:
        idx = find_paragraph_index(doc, cand)
        if idx >= 0:
            next_anchor = cand
            break
    if next_anchor is None:
        raise RuntimeError("Could not locate §5.5 anchor for §5.4.3 boundary")
    print(f"§5.5 anchor identified as: '{next_anchor}'")

    print("Replacing §5.4.3 content...")
    replace_section_content(
        doc, "5.4.3 Collision Avoidance", next_anchor,
        list(reversed(CONTENT_5_4_3)),
    )

    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!\n")
    print("Summary of changes (all yellow-highlighted):")
    print("  §5.4 intro: rewritten to name the three components and the 2.0-second cycle")
    print("  §5.4.1: full scoring formula, weights explained, cap at max_targets=200")
    print("  §5.4.2: corrected 'greedy' → 'Hungarian (linear_sum_assignment)';")
    print("           full scoring formula; hysteresis 0.1; RTB/CHARGING exclusion")
    print("  §5.4.3: VO time-to-closest-approach formula; KDTree query; 50/50 responsibility;")
    print("           firefighter exclusion 15 m / strength 3.0; clarified that fire standoff")
    print("           is in _safe_hover_pos geometry, not the collision avoider")


if __name__ == "__main__":
    main()
