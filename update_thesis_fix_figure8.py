#!/usr/bin/env python3
"""Replace Figure 8 in §5.9.3 with the current 35,000-tick total-active-fire figure.

The original Figure 8 was inserted by an earlier update script and shows the
old 5,000-tick (500 simulated seconds) acoustic-only fleet-size sweep. This
script replaces both the image and the caption with the current long-run
data covering the full 35,000-tick (3,500 simulated seconds) window across the
mechanism-comparison scenarios.

Builds on Thesis_04_29_Section5_v7.docx. Saves as new file.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

INPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_v7.docx"
OUTPUT_PATH = "/Users/xtin2000/Documents/LMU/Thesis/Thesis_04_29_Section5_v8.docx"
FIG_PATH = "/Users/xtin2000/wildfire_drone_swarm/results_v2/figures/fig_active_fire.png"

NEW_CAPTION = (
    "Figure 8: Total active fire cells (burning + ember) over simulated time across the "
    "current mechanism-comparison scenarios at 5 m/s wind (combined_high_wind shown dashed at "
    "10 m/s). Each scenario was run for the full 35,000-tick (3,500 simulated second, "
    "≈ 58-minute) window or until mission_complete, replacing the 5,000-tick (500 simulated "
    "second) windows used in the earlier acoustic-only fleet-size sweep (§5.6.1). The "
    "combined_baseline configuration (green) collapses to zero within the first 17 simulated "
    "minutes; all other configurations grow into the thousands of active fire cells and run "
    "the full window without achieving containment."
)


def find_paragraph_by_caption(doc, caption_prefix: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith(caption_prefix):
            return i
    return -1


def main() -> None:
    print(f"Loading: {INPUT_PATH}")
    if not os.path.exists(INPUT_PATH):
        raise FileNotFoundError(INPUT_PATH)
    doc = Document(INPUT_PATH)

    # Find the existing Figure 8 caption paragraph
    cap_idx = find_paragraph_by_caption(doc, "Figure 8: Total active fire cells")
    if cap_idx < 0:
        raise RuntimeError("Could not find existing 'Figure 8: Total active fire cells' caption")
    print(f"Found Figure 8 caption at paragraph {cap_idx}")

    # The image is in the paragraph immediately before the caption
    fig_idx = cap_idx - 1
    fig_para = doc.paragraphs[fig_idx]
    cap_para = doc.paragraphs[cap_idx]

    # Anchor: paragraph just before the figure (we'll insert the new figure+caption after it)
    anchor = doc.paragraphs[fig_idx - 1]

    # Remove the old caption first, then the old figure
    cap_para._element.getparent().remove(cap_para._element)
    fig_para._element.getparent().remove(fig_para._element)

    # Insert new caption first, then new figure (addnext, so reverse-order insertion)
    new_cap_p = OxmlElement("w:p")
    anchor._element.addnext(new_cap_p)
    for p in doc.paragraphs:
        if p._element is new_cap_p:
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(NEW_CAPTION)
            run.italic = True
            run.font.size = Pt(10)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            break

    new_fig_p = OxmlElement("w:p")
    anchor._element.addnext(new_fig_p)
    for p in doc.paragraphs:
        if p._element is new_fig_p:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(FIG_PATH, width=Inches(5.5))
            break

    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done!\n")
    print("Summary:")
    print("  - Replaced Figure 8 image (was 5,000-tick acoustic-only fleet-size sweep)")
    print("    with the current 35,000-tick mechanism-comparison total-active-fire figure")
    print("  - Updated caption to reflect the long-run window and the current scenarios")


if __name__ == "__main__":
    main()
